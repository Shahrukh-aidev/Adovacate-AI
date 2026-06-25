import gradio as gr
import json
import os
import re
from datetime import datetime
import requests
from pinecone import Pinecone
import cohere

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY

from docx import Document as DocxDocument
from docx.shared import Pt, Cm as DocxCm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ========== DOCUMENT UPLOAD & ANALYSIS (OCR + LLM) ==========
try:
    import pytesseract
    from PIL import Image
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    print("⚠️ pytesseract/PIL not installed. OCR disabled. Install: pip install pytesseract Pillow")

try:
    import pdf2image
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    print("⚠️ pdf2image not installed. PDF OCR disabled. Install: pip install pdf2image")

# NEW: pdfplumber for native PDF text extraction (much faster & better for text-based PDFs)
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    print("⚠️ pdfplumber not installed. Native PDF text extraction disabled. Install: pip install pdfplumber")

def extract_text_from_image(image_path):
    if not TESSERACT_AVAILABLE or not image_path:
        return None
    try:
        img = Image.open(image_path)
        if img.mode != 'RGB':
            img = img.convert('RGB')
        text = pytesseract.image_to_string(img, lang='eng+urd')
        if not text.strip():
            text = pytesseract.image_to_string(img, lang='eng')
        return text.strip()
    except Exception as e:
        print(f"OCR image error: {e}")
        return None

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF using native extraction (pdfplumber) first, fallback to OCR."""
    # 1. Try native text extraction (for digital PDFs)
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
                if text.strip():
                    return text.strip()
        except Exception as e:
            print(f"pdfplumber extraction failed: {e}")

    # 2. Fallback to OCR (for scanned PDFs)
    if not PDF2IMAGE_AVAILABLE or not TESSERACT_AVAILABLE:
        return None
    try:
        images = pdf2image.convert_from_path(pdf_path, dpi=200)
        all_text = []
        for i, img in enumerate(images):
            text = pytesseract.image_to_string(img, lang='eng+urd')
            if not text.strip():
                text = pytesseract.image_to_string(img, lang='eng')
            all_text.append(f"--- Page {i+1} ---\n{text.strip()}")
        return "\n\n".join(all_text)
    except Exception as e:
        print(f"OCR PDF error: {e}")
        return None

def analyze_document(doc_text, doc_type):
    if not doc_text or len(doc_text.strip()) < 20:
        return "⚠️ Could not extract readable text from the document. Please ensure the file is clear and try again."
    doc_text = doc_text[:4000]
    prompt = f"""You are HAQ, Pakistan's AI legal assistant. Analyze the following {doc_type} and provide a comprehensive legal breakdown.
DOCUMENT TYPE: {doc_type}
EXTRACTED TEXT:
{doc_text}
Provide your analysis in this EXACT format:
📄 DOCUMENT OVERVIEW
[2-3 sentence summary of what this document is]
⚠️ CRITICAL CLAUSES / WARNINGS
1. [Clause 1]: [What it means in simple terms]
2. [Clause 2]: [What it means in simple terms]
3. [Clause 3]: [What it means in simple terms]
📋 KEY DATES & DEADLINES
• [Date/Deadline 1]: [What action is required]
• [Date/Deadline 2]: [What action is required]
⚖️ LEGAL IMPLICATIONS
[Explain the legal consequences and what laws apply]
✅ WHAT YOU SHOULD DO NEXT
1. [Immediate step]
2. [Short-term step]
3. [Long-term step]
🚨 RED FLAGS (if any)
[Any suspicious, unfair, or legally problematic clauses]
📞 WHERE TO GET HELP
[Specific court, lawyer type, or government office to contact]
DISCLAIMER: This is general legal information. Consult a licensed Vakeel (lawyer) for official advice.
RULES:
- Use simple language — the user may not be a lawyer
- Cite specific Pakistani laws where relevant
- If it's a court notice, explain what will happen if they don't respond
- If it's a contract, highlight one-sided or risky clauses
- If it's an FIR, explain the charges and next steps
- If it's a legal notice, explain the deadline and consequences
- Be honest if the document text is unclear or incomplete"""
    msgs = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "user", "content": prompt}
    ]
    answer, err = call_groq(msgs, max_tokens=1800)
    if answer:
        return answer
    return f"⚠️ Analysis failed: {err}"

def process_uploaded_document(file_path, doc_type):
    if not file_path:
        return "Please upload a document first.", None
    file_path = str(file_path)
    ext = os.path.splitext(file_path)[1].lower()
    extracted_text = None
    if ext in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp'):
        extracted_text = extract_text_from_image(file_path)
    elif ext == '.pdf':
        extracted_text = extract_text_from_pdf(file_path)
    elif ext in ('.txt', '.md'):
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f2:
                extracted_text = f2.read()
        except Exception as e:
            return f"⚠️ Could not read text file: {e}", None
    else:
        return f"⚠️ Unsupported file type: {ext}. Please upload PDF, image (PNG/JPG), or text file.", None
    if not extracted_text:
        return "⚠️ Could not extract text from the document. The file may be corrupted, scanned without OCR, or in an unsupported format.", None
    analysis = analyze_document(extracted_text, doc_type)
    preview = f"📄 EXTRACTED TEXT PREVIEW (first 800 chars):\n{'─'*60}\n{extracted_text[:800]}..."
    return analysis, preview

# ========== TEXT-TO-SPEECH (VOICE OUTPUT) ==========
# Improved TTS: gTTS with better cleaning, plus optional Coqui TTS if installed.

try:
    from gtts import gTTS
    GTTS_AVAILABLE = True
except ImportError:
    GTTS_AVAILABLE = False
    print("⚠️ gTTS not installed. Voice output disabled. Install: pip install gTTS")

# Optional: better TTS (Coqui TTS) – uncomment if you install it
# try:
#     from TTS.api import TTS
#     COQUI_AVAILABLE = True
#     tts = TTS("tts_models/ur/fairseq/vits")  # Urdu model
# except:
#     COQUI_AVAILABLE = False
#     tts = None

def detect_language(text):
    urdu_chars = sum(1 for c in text if '\u0600' <= c <= '\u06FF' or '\u0750' <= c <= '\u077F')
    total_chars = sum(1 for c in text if c.isalpha())
    if total_chars == 0:
        return 'en'
    return 'ur' if urdu_chars / total_chars > 0.3 else 'en'

def text_to_speech(answer_text):
    if not GTTS_AVAILABLE or not answer_text or len(answer_text.strip()) < 3:
        return None
    try:
        # Better cleaning for TTS: remove markdown, citations, etc.
        clean_text = re.sub(r'\[.*?\]', '', answer_text)          # remove [citations]
        clean_text = re.sub(r'\([^)]*\)', '', clean_text)         # remove (parentheticals) – optional
        clean_text = re.sub(r'[#*`_~\[\](){}>]', '', clean_text)  # remove markdown symbols
        clean_text = re.sub(r'<[^>]+>', '', clean_text)           # remove HTML tags
        clean_text = re.sub(r'[^\w\s\.\,\!\?\u0600-\u06FF\u0750-\u077F]', ' ', clean_text)  # keep Urdu + punctuation
        clean_text = re.sub(r'\s+', ' ', clean_text).strip()[:500] # limit length
        if not clean_text:
            return None

        lang = detect_language(clean_text)
        # Use gTTS with standard speed (slow=False) and TLD='com' for slightly better voice
        tts = gTTS(text=clean_text, lang=lang, slow=False, tld='com')
        audio_path = f"/tmp/haq_tts_{datetime.now().strftime('%Y%m%d_%H%M%S')}.mp3"
        tts.save(audio_path)
        return audio_path

    except Exception as e:
        print(f"TTS error: {e}")
        return None

# ========== VOICE INPUT FUNCTIONALITY ==========
def transcribe_voice(audio_path):
    if not audio_path or not GROQ_KEY:
        return "⚠️ Audio recording failed or GROQ_KEY missing. Please type your question."
    try:
        with open(audio_path, "rb") as audio_file:
            r = requests.post(
                "https://api.groq.com/openai/v1/audio/transcriptions",
                headers={"Authorization": f"Bearer {GROQ_KEY}"},
                files={"file": ("voice.wav", audio_file, "audio/wav")},
                data={"model": "whisper-large-v3", "language": "ur", "response_format": "json"},
                timeout=30
            )
        if r.status_code == 200:
            text = r.json().get("text", "").strip()
            if text:
                return text
            return "⚠️ Could not understand audio. Please speak clearly or type your question."
        else:
            return f"⚠️ Transcription error (HTTP {r.status_code}). Please type your question."
    except Exception as e:
        print(f"Transcription error: {e}")
        return "⚠️ Voice transcription failed. Please type your question."

# ========== API KEYS & CONNECTIONS ==========
PINECONE_KEY = os.environ.get("PINECONE_KEY", "").strip()
COHERE_KEY   = os.environ.get("COHERE_KEY", "").strip()
GROQ_KEY     = os.environ.get("GROQ_KEY", "").strip()

print(f"Pinecone: {'OK' if PINECONE_KEY else 'MISSING'}")
print(f"Cohere:   {'OK' if COHERE_KEY else 'MISSING'}")
print(f"Groq:     {'OK' if GROQ_KEY else 'MISSING'}")

index = None
co    = None

try:
    if PINECONE_KEY:
        pc = Pinecone(api_key=PINECONE_KEY)
        index = pc.Index("haq-law")
        print("Pinecone connected")
except Exception as e:
    print(f"Pinecone error: {e}")

try:
    if COHERE_KEY:
        co = cohere.Client(COHERE_KEY)
        print("Cohere connected")
except Exception as e:
    print(f"Cohere error: {e}")

# ========== LAW LINKS & PATTERNS ==========
LAW_LINKS = {
    "constitution": {"name": "Constitution of Pakistan 1973", "url": "https://pakistancode.gov.pk/english/UY2FqaJw1-apaUY2Fqa-ap8=-sg-jjjjjjjjjjjjj", "icon": "🏛"},
    "ppc": {"name": "Pakistan Penal Code 1860", "url": "https://pakistancode.gov.pk/english/UY2FqaJw1-apaUY2Fqa-apaUY2Npa5lo-sg-jjjjjjjjjjjjj", "icon": "⚖"},
    "crpc": {"name": "Code of Criminal Procedure 1898", "url": "https://pakistancode.gov.pk/english/UY2FqaJw1-apaUY2Fqa-apaUY2Npa5lp-sg-jjjjjjjjjjjjj", "icon": "📋"},
    "peca": {"name": "Prevention of Electronic Crimes Act 2016", "url": "https://pakistancode.gov.pk/english/UY2FqaJw1-apaUY2Fqa-apaUY2Jvbp8=-sg-jjjjjjjjjjjjj", "icon": "💻"},
    "mflo": {"name": "Muslim Family Laws Ordinance 1961", "url": "https://pakistancode.gov.pk/english/UY2FqaJw1-apaUY2Fqa-apaUY2Npa5po-sg-jjjjjjjjjjjjj", "icon": "👪"},
    "contract": {"name": "Contract Act 1872", "url": "https://pakistancode.gov.pk/english/UY2FqaJw2-apaUY2Fqa-a50=-sg-jjjjjjjjjjjjj-con-51", "icon": "📝"},
    "property": {"name": "Transfer of Property Act 1882", "url": "https://pakistancode.gov.pk/english/UY2FqaJw1-apaUY2Fqa-bpk=-sg-jjjjjjjjjjjjj", "icon": "🏠"},
    "labour": {"name": "Payment of Wages Act 1936", "url": "https://pakistancode.gov.pk/english/UY2FqaJw1-apaUY2Fqa-cJY=-sg-jjjjjjjjjjjjj", "icon": "💼"},
    "registration": {"name": "Registration Act 1908", "url": "https://pakistancode.gov.pk/english/UY2FqaJw1-apaUY2Fqa-apeU-sg-jjjjjjjjjjjjj", "icon": "📄"},
    "specific_relief": {"name": "Specific Relief Act 1877", "url": "https://pakistancode.gov.pk/english/UY2FqaJw1-apaUY2Fqa-bpo=-sg-jjjjjjjjjjjjj", "icon": "🔍"},
    "ata": {"name": "Anti-Terrorism Act 1997", "url": "https://pakistancode.gov.pk/english/UY2FqaJw1-apaUY2Fqa-apaUY2Npappq-sg-jjjjjjjjjjjjj", "icon": "🚨"},
    "dmma": {"name": "Dissolution of Muslim Marriages Act 1939", "url": "https://pakistancode.gov.pk/english/UY2FqaJw1-apaUY2Fqa-cJaW-sg-jjjjjjjjjjjjj", "icon": "📜"},
    "evidence": {"name": "Qanoon-e-Shahadat Order 1984", "url": "https://pakistancode.gov.pk/english/UY2FqaJw1-apaUY2Fqa-apaUY2Npa5plaw==-sg-jjjjjjjjjjjjj", "icon": "🔎"},
    "banking": {"name": "Banking Companies Ordinance 1962", "url": "https://pakistancode.gov.pk/english/UY2FqaJw1-apaUY2Fqa-apaUY2Npa5qq-sg-jjjjjjjjjjjjj", "icon": "🏦"},
}

LAW_PATTERNS = {
    "constitution": [r'[Aa]rticle\s*\d+', r'[Cc]onstitution', r'[Ff]undamental [Rr]ight', r'[Hh]abeas [Cc]orpus', r'[Mm]andamus', r'[Ww]rit'],
    "ppc": [r'PPC', r'[Pp]enal [Cc]ode', r'[Ss]ection\s*30[0-9]', r'[Ss]ection\s*37[0-9]', r'[Ss]ection\s*4[12][0-9]', r'[Qq]atl', r'[Dd]iyat', r'[Rr]ape', r'[Tt]heft', r'[Cc]heating'],
    "crpc": [r'CrPC', r'[Cc]riminal [Pp]rocedure', r'[Ss]ection\s*154', r'[Ss]ection\s*497', r'[Ss]ection\s*498', r'[Ss]ection\s*54\b', r'[Ff]IR', r'[Bb]ail\b', r'[Cc]hallan'],
    "peca": [r'PECA', r'[Cc]ybercrime', r'[Ee]lectronic [Cc]rime', r'[Oo]nline [Hh]arassment', r'FIA\b', r'[Cc]yberstalking'],
    "mflo": [r'MFLO', r'[Mm]uslim [Ff]amily', r'[Kk]hula', r'[Tt]alaq', r'[Nn]ikah', r'[Mm]ehr', r'[Ii]ddat'],
    "contract": [r'[Cc]ontract [Aa]ct', r'[Ss]ection\s*73\b', r'[Ss]ection\s*19\b', r'[Bb]reach of [Cc]ontract', r'[Cc]onsideration'],
    "property": [r'[Tt]ransfer of [Pp]roperty', r'[Ss]ection\s*53.?[Aa]', r'TPA\b', r'[Mm]ortgage', r'[Aa]dverse [Pp]ossession'],
    "labour": [r'[Pp]ayment of [Ww]ages', r'[Ll]abour [Cc]ourt', r'[Mm]inimum [Ww]ages', r'[Gg]ratuity', r'EOBI', r'[Ww]orkmen [Cc]ompensation'],
    "registration": [r'[Rr]egistration [Aa]ct', r'[Ss]ub-[Rr]egistrar', r'[Ss]tamp [Dd]uty'],
    "specific_relief": [r'[Ss]pecific [Rr]elief', r'[Ss]pecific [Pp]erformance'],
    "ata": [r'[Aa]nti-[Tt]errorism', r'\bATA\b', r'\bATC\b'],
    "dmma": [r'DMMA', r'[Dd]issolution of [Mm]uslim'],
    "evidence": [r'[Qq]anoon-e-[Ss]hahadat', r'[Cc]onfession', r'[Ee]lectronic [Ee]vidence'],
    "banking": [r'[Bb]anking [Cc]ompanies', r'[Bb]anking [Mm]ohtasib', r'SBP\b'],
}

def detect_laws(text):
    found = []
    for law_key, patterns in LAW_PATTERNS.items():
        for pat in patterns:
            if re.search(pat, text):
                found.append(law_key)
                break
    return found

def build_verify_html(answer_text):
    detected = detect_laws(answer_text)
    if not detected:
        return ""
    cards = ""
    for law_key in detected:
        law = LAW_LINKS[law_key]
        cards += (
            '<a href="' + law['url'] + '" target="_blank" rel="noopener noreferrer"'
            ' style="display:flex;align-items:center;gap:12px;padding:11px 14px;'
            'background:rgba(255,255,255,0.03);border:1px solid rgba(255,255,255,0.06);'
            'border-radius:10px;text-decoration:none;margin-bottom:7px;">'
            '<span style="font-size:20px;flex-shrink:0;">' + law['icon'] + '</span>'
            '<div style="flex:1;min-width:0;">'
            '<div style="color:#e2e8f0;font-size:13px;font-weight:600;">' + law['name'] + '</div>'
            '<div style="color:#64748b;font-size:11px;margin-top:2px;">pakistancode.gov.pk — Ministry of Law &amp; Justice</div>'
            '</div>'
            '<div style="background:rgba(16,185,129,0.15);border:1px solid rgba(16,185,129,0.3);'
            'color:#10b981;font-size:11px;font-weight:700;padding:4px 10px;'
            'border-radius:20px;flex-shrink:0;">Verify</div>'
            '</a>'
        )
    return (
        '<div style="margin:12px 0 0;padding:16px;background:rgba(16,185,129,0.05);'
        'border:1px solid rgba(16,185,129,0.2);border-radius:14px;">'
        '<div style="color:#10b981;font-size:12px;font-weight:700;text-transform:uppercase;'
        'letter-spacing:0.8px;margin-bottom:12px;">'
        '✓ Verify These Laws — Official Sources'
        '</div>'
        + cards +
        '<div style="color:#475569;font-size:11px;margin-top:8px;text-align:center;'
        'padding-top:8px;border-top:1px solid rgba(255,255,255,0.05);">'
        'All links → pakistancode.gov.pk — Pakistan Ministry of Law &amp; Justice (Official)'
        '</div></div>'
    )

SYSTEM_PROMPT = """You are HAQ, Pakistan's most accurate AI legal assistant. You remember the full conversation and answer follow-up questions in context.
CORE RULES:
1. ALWAYS remember what was discussed earlier in this conversation
2. When user asks follow-up like "what if he doesn't comply?" or "phir kya hoga?" — refer back to the previous legal situation
3. Cite exact law name + section number always
4. Never invent citations. Never give illegal advice.
5. Urdu question -> Pure Urdu answer. English -> English. Mixed -> Match user's language.
6. NEVER use Hindi words — pure Urdu only (afsar not adhikari, karwai not karyavahi)
7. Keep answers focused — don't repeat information already given in this conversation
FORMAT EVERY ANSWER:
LEGAL BASIS
[Law Name Year, Section X]: [explanation]
THE RULING
[Direct answer in 2-3 sentences]
WHAT YOU SHOULD DO
1. [Step 1]
2. [Step 2]
3. [Step 3]
WHERE TO GO
[Specific authority/court]
DISCLAIMER
General legal information. Consult a licensed Vakeel for court cases.
KEY LAWS:
Constitution 1973: Art 9 (life/liberty), Art 10 (arrest safeguards-24hrs-lawyer), Art 10-A (fair trial), Art 13 (double jeopardy), Art 25 (equality), Art 199 (High Court writs)
PPC 1860: Sec 302 (qatl-i-amd-death/life), Sec 375-376 (rape), Sec 379 (theft-3yrs), Sec 420 (cheating-7yrs), Sec 441 (trespass), Sec 489-F (bad cheque-3yrs), Sec 499-500 (defamation-2yrs)
CrPC 1898: Sec 22A (Justice of Peace), Sec 54 (arrest without warrant), Sec 154 (FIR-mandatory-free), Sec 497 (bail), Sec 498 (anticipatory bail-High Court), Sec 561A (FIR quashment)
MFLO 1961: Sec 6 (polygamy-Arbitration Council), Sec 7 (talaq-90days), Sec 8 (khula-no husband consent needed)
PECA 2016: Sec 20 (harassment-3yrs/Rs1M), Sec 21 (private images-5yrs/Rs5M), Sec 24 (cyberstalking-3yrs)
Labour: Wages by 7th, show-cause before termination, gratuity=30days/yr after 5yrs, overtime=2x
Contract Act 1872: Sec 10 (valid contract), Sec 19 (voidable), Sec 73 (breach compensation)
Property: TPA 1882 Sec 53A, adverse possession=12yrs
Evidence: Qanoon-e-Shahadat 1984 Art 35 (police confession inadmissible), Art 164 (electronic evidence)"""

LETTER_SYSTEM_PROMPT = """You are HAQ, Pakistan's AI legal assistant. Generate professional Pakistani legal notice body text.
Write EXACTLY 6 paragraphs. Do NOT write greeting, salutation, header, closing, or signature.
PARAGRAPH 1 — INTRODUCTION: Who sender is and general nature of grievance.
PARAGRAPH 2 — STATEMENT OF FACTS: Chronological facts with dates/amounts where provided.
PARAGRAPH 3 — LEGAL VIOLATIONS: "Your actions constitute a violation of [Law Name Year], Section [X], which provides that [brief text]."
PARAGRAPH 4 — TAKE NOTICE: "TAKE NOTICE THAT you are hereby called upon and required to [specific demand] within FIFTEEN (15) days from receipt of this notice."
PARAGRAPH 5 — CONSEQUENCES: Specific court + sections that will be invoked upon non-compliance.
PARAGRAPH 6 — COSTS: Liability for all costs and damages.
RULES: Formal Pakistani legal English. Cite exact law + section for every claim. 15-day deadline standard."""

def get_embedding(text):
    if not co:
        return None
    try:
        r = co.embed(texts=[str(text)[:500]], model="embed-english-light-v3.0", input_type="search_query")
        return r.embeddings[0]
    except Exception as e:
        print(f"Embed error: {e}")
        return None

def search_laws(question, top_k=5):
    if not index:
        return []
    emb = get_embedding(question)
    if not emb:
        return []
    try:
        r = index.query(vector=emb, top_k=top_k, include_metadata=True)
        return [{'law': m.metadata.get('law_name','?'),
                 'text': str(m.metadata.get('text',''))[:400],
                 'score': round(m.score, 3)}
                for m in r.matches if m.score > 0.2]
    except Exception as e:
        print(f"Search error: {e}")
        return []

def call_groq(messages, max_tokens=1400):
    if not GROQ_KEY:
        return None, "No GROQ_KEY — add it in HuggingFace Space Secrets"
    models = [
        "llama-3.1-8b-instant",
        "llama-3.3-70b-versatile",
        "gemma2-9b-it",
        "llama3-70b-8192",
    ]
    last_error = "Unknown"
    for model in models:
        try:
            print(f"Trying: {model}")
            r = requests.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers={"Authorization": f"Bearer {GROQ_KEY}", "Content-Type": "application/json"},
                json={"model": model, "temperature": 0.1, "max_tokens": max_tokens, "messages": messages},
                timeout=45
            )
            if r.status_code == 200:
                content = r.json().get('choices', [{}])[0].get('message', {}).get('content', '')
                if content and len(content) > 30:
                    print(f"OK: {model}")
                    return content, None
            elif r.status_code == 401:
                return None, "Invalid GROQ_KEY"
            else:
                last_error = f"HTTP {r.status_code} on {model}"
        except requests.exceptions.Timeout:
            last_error = f"Timeout on {model}"
        except Exception as e:
            last_error = str(e)[:60]
    return None, f"All models failed. Last: {last_error}"

def chat_with_haq(user_message, history):
    if not user_message or not user_message.strip():
        history.append([user_message, "Please write your legal question."])
        return history, ""
    user_message = user_message.strip()
    for p in ['ignore all laws', 'ignore all instructions', 'help me commit', 'how to kill']:
        if p in user_message.lower():
            history.append([user_message, "HAQ only answers genuine legal questions."])
            return history, ""
    sections = search_laws(user_message)
    ctx = ""
    if sections:
        ctx = "\nRELEVANT LAW SECTIONS FROM DATABASE:\n"
        for i, s in enumerate(sections, 1):
            ctx += f"[{i}] {s['law']} (score:{s['score']})\n{s['text']}\n"
    messages = [{"role": "system", "content": SYSTEM_PROMPT}]
    for user_turn, bot_turn in history:
        if user_turn:
            messages.append({"role": "user", "content": str(user_turn)})
        if bot_turn:
            messages.append({"role": "assistant", "content": str(bot_turn)})
    current_content = f"{ctx}\nQUESTION: {user_message}" if ctx else f"QUESTION: {user_message}"
    messages.append({"role": "user", "content": current_content})
    if len(messages) > 22:
        messages = [messages[0]] + messages[-21:]
    answer, err = call_groq(messages, max_tokens=1200)
    if not answer:
        answer = f"⚠️ Service temporarily unavailable: {err}\nCheck GROQ_KEY in HuggingFace Secrets."
    history.append([user_message, answer])
    return history, ""

def get_rights(situation):
    if not situation:
        return "Please select your situation.", ""
    question = f"What are all my legal rights in Pakistan for: {situation}. Cite exact law + section, step-by-step advice, which authority to contact."
    sections = search_laws(question)
    ctx = "RELEVANT LAW SECTIONS:\n"
    if sections:
        for i, s in enumerate(sections, 1):
            ctx += f"[{i}] {s['law']}\n{s['text']}\n"
    msgs = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "user",   "content": f"{ctx}\nQUESTION: {question}"}
    ]
    answer, err = call_groq(msgs)
    if answer:
        return answer, build_verify_html(answer)
    return f"⚠️ {err}", ""

def generate_pdf_file(letter_body, your_name, other_name, your_address, letter_type):
    try:
        path = f"/tmp/HAQ_Legal_Notice_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        doc = SimpleDocTemplate(path, pagesize=A4, rightMargin=2.5*cm, leftMargin=2.5*cm, topMargin=2*cm, bottomMargin=2*cm)
        story = []
        s = getSampleStyleSheet()
        def st(name, **kw):
            base = kw.pop('parent', s['Normal'])
            return ParagraphStyle(name, parent=base, **kw)
        S_HDR = st('H', fontSize=16, fontName='Helvetica-Bold', alignment=TA_CENTER, textColor=colors.white)
        S_SUB = st('Su', fontSize=9, fontName='Helvetica-Oblique', alignment=TA_CENTER, textColor=colors.HexColor('#666'), spaceAfter=4)
        S_REF = st('Rf', fontSize=9, fontName='Helvetica', textColor=colors.HexColor('#555'))
        S_REFR = st('Rr', fontSize=9, fontName='Helvetica', alignment=TA_RIGHT, textColor=colors.HexColor('#555'))
        S_PL = st('PL', fontSize=10, fontName='Helvetica', textColor=colors.HexColor('#1a1a2e'), leading=17)
        S_PR = st('PR', fontSize=10, fontName='Helvetica', alignment=TA_RIGHT, textColor=colors.HexColor('#8B0000'), leading=17)
        S_MODE = st('MD', fontSize=8, fontName='Helvetica-Bold', alignment=TA_CENTER, textColor=colors.HexColor('#555'), spaceAfter=8)
        S_SUBJ = st('SJ', fontSize=10, fontName='Helvetica-Bold', alignment=TA_CENTER, textColor=colors.HexColor('#8B0000'), spaceAfter=8)
        S_BOLD = st('BB', fontSize=10, fontName='Helvetica-Bold', spaceAfter=4)
        S_BODY = st('B', fontSize=10, fontName='Helvetica', alignment=TA_JUSTIFY, spaceAfter=9, leading=17, textColor=colors.HexColor('#1a1a1a'))
        S_SIGN = st('SG', fontSize=10, fontName='Helvetica', leading=17)
        S_SIGNR = st('SR', fontSize=10, fontName='Helvetica', alignment=TA_RIGHT, leading=17)
        S_ACKH = st('AH', fontSize=9, fontName='Helvetica-Bold', alignment=TA_CENTER, textColor=colors.HexColor('#333'))
        S_ACK = st('AK', fontSize=9, fontName='Helvetica', textColor=colors.HexColor('#444'), leading=15)
        S_FOOT = st('FT', fontSize=7.5, fontName='Helvetica', alignment=TA_CENTER, textColor=colors.HexColor('#888'))
        date_str = datetime.now().strftime("%d %B %Y")
        ref_num = f"LN/{datetime.now().year}/{datetime.now().strftime('%m%d%H%M')}"
        hdr = Table([[Paragraph("⚖   LEGAL NOTICE   ⚖", S_HDR)]], colWidths=[16*cm])
        hdr.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#1a1a2e')),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('TOPPADDING', (0,0), (-1,-1), 14), ('BOTTOMPADDING', (0,0), (-1,-1), 14),
        ]))
        story += [hdr, Spacer(1,4)]
        story.append(Paragraph("Under the Laws of the Islamic Republic of Pakistan", S_SUB))
        story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor('#C9A84C')))
        story.append(Spacer(1,8))
        story.append(Table([
            [Paragraph(f"Ref. No: <b>{ref_num}</b>", S_REF), Paragraph(f"Date: <b>{date_str}</b>", S_REFR)]
        ], colWidths=[8*cm, 8*cm]))
        story.append(Spacer(1,10))
        p = Table([[Paragraph(f"<b>FROM (Murasil / Noticee):</b><br/><b>{your_name or 'Applicant'}</b><br/>{your_address or 'Pakistan'}", S_PL),
                    Paragraph(f"<b>TO (Mukhatib / Respondent):</b><br/><b>{other_name or 'Respondent'}</b><br/>Address as Known to Sender", S_PR)
        ]], colWidths=[7.8*cm, 8.2*cm])
        p.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (0,0), colors.HexColor('#eef2ff')),
            ('BACKGROUND', (1,0), (1,0), colors.HexColor('#fff0f0')),
            ('BOX', (0,0), (0,0), 0.5, colors.HexColor('#9ba8d0')),
            ('BOX', (1,0), (1,0), 0.5, colors.HexColor('#d0a0a0')),
            ('TOPPADDING', (0,0), (-1,-1), 10), ('BOTTOMPADDING', (0,0), (-1,-1), 10),
            ('LEFTPADDING', (0,0), (-1,-1), 10), ('RIGHTPADDING', (0,0), (-1,-1), 10),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ]))
        story += [p, Spacer(1,8)]
        story.append(Paragraph("MODE OF SERVICE: REGISTERED POST A.D. / IN PERSON / COURIER", S_MODE))
        story.append(HRFlowable(width="100%", thickness=2.5, color=colors.HexColor('#1a1a2e')))
        story.append(Spacer(1,10))
        sb = Table([[Paragraph(f"SUBJECT: {letter_type.upper()}", S_SUBJ)]], colWidths=[16*cm])
        sb.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#fff5f5')),
            ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#e8c0c0')),
            ('TOPPADDING', (0,0), (-1,-1), 8), ('BOTTOMPADDING', (0,0), (-1,-1), 8),
        ]))
        story += [sb, Spacer(1,10)]
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#cccccc')))
        story += [Spacer(1,8), Paragraph("Sir / Madam,", S_BOLD), Spacer(1,6)]
        skip = ('sir/', 'madam', 'sincerely', 'yours faithfully', 'yours truly', 'yours sincerely', 'respectfully,', 'regards,')
        for para in [p.strip() for p in letter_body.split('\n') if p.strip()]:
            if any(para.lower().startswith(x) for x in skip):
                continue
            story.append(Paragraph(para, S_BODY))
        story += [Spacer(1,14), HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#cccccc')), Spacer(1,10)]
        story.append(Paragraph("Yours faithfully,", S_BODY))
        story.append(Spacer(1,28))
        story.append(Table([[Paragraph(f"________________________<br/><b>{your_name or 'Applicant'}</b><br/><i>Noticee / Applicant</i>", S_SIGN),
                              Paragraph("________________________<br/><b>Advocate / Legal Counsel</b><br/><i>(If applicable)</i>", S_SIGNR)
        ]], colWidths=[8*cm, 8*cm]))
        story.append(Spacer(1,20))
        ack_h = Table([[Paragraph("✦  ACKNOWLEDGEMENT OF RECEIPT  ✦", S_ACKH)]], colWidths=[16*cm])
        ack_h.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f5f5f5')),
            ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#aaaaaa')),
            ('TOPPADDING', (0,0), (-1,-1), 7), ('BOTTOMPADDING', (0,0), (-1,-1), 7),
        ]))
        story.append(ack_h)
        ack_b = Table([[Paragraph(
            f"I, <b>{other_name or 'Respondent'}</b>, hereby acknowledge receipt of this Legal Notice on "
            f"_____________ at _____________.<br/><br/>"
            f"Signature: _________________________   Date: _________________________<br/><br/>"
            f"Witness: _________________________", S_ACK)
        ]], colWidths=[16*cm])
        ack_b.setStyle(TableStyle([
            ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#aaaaaa')),
            ('TOPPADDING', (0,0), (-1,-1), 10), ('BOTTOMPADDING', (0,0), (-1,-1), 10),
            ('LEFTPADDING', (0,0), (-1,-1), 12), ('RIGHTPADDING', (0,0), (-1,-1), 12),
        ]))
        story += [ack_b, Spacer(1,12)]
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#cccccc')))
        story.append(Paragraph(
            f"Generated by HAQ — Pakistan Legal AI  |  {date_str}  |  Ref: {ref_num}  |  "
            "General information only. Consult a licensed Vakeel for court matters.", S_FOOT))
        doc.build(story)
        return path
    except Exception as e:
        print(f"PDF error: {e}")
        return None

def generate_docx_file(letter_body, your_name, other_name, your_address, letter_type):
    try:
        path = f"/tmp/HAQ_Legal_Notice_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc = DocxDocument()
        sec = doc.sections[0]
        sec.page_width = DocxCm(21); sec.page_height = DocxCm(29.7)
        sec.left_margin = DocxCm(2.8); sec.right_margin = DocxCm(2.8)
        sec.top_margin = DocxCm(2); sec.bottom_margin = DocxCm(2)
        date_str = datetime.now().strftime("%d %B %Y")
        ref_num = f"LN/{datetime.now().year}/{datetime.now().strftime('%m%d%H%M')}"
        def add_para(text='', bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_before=0, space_after=6, font_name='Times New Roman'):
            p = doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after = Pt(space_after)
            if text:
                r = p.add_run(text)
                r.bold = bold; r.italic = italic
                r.font.name = font_name; r.font.size = Pt(size)
                if color: r.font.color.rgb = RGBColor(*color)
            return p
        def shade_para(para, hex_color):
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
            para._p.get_or_add_pPr().append(shd)
        def add_border(para, side='bottom', sz=6, color='000000'):
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bd = OxmlElement(f'w:{side}')
            bd.set(qn('w:val'), 'single'); bd.set(qn('w:sz'), str(sz))
            bd.set(qn('w:space'), '1'); bd.set(qn('w:color'), color)
            pBdr.append(bd); pPr.append(pBdr)
        def no_border(cell):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcBdr = OxmlElement('w:tcBorders')
            for side in ('top','left','bottom','right','insideH','insideV'):
                bd = OxmlElement(f'w:{side}'); bd.set(qn('w:val'), 'none'); tcBdr.append(bd)
            tcPr.append(tcBdr)
        def cell_fill(cell, fill):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
        hdr = add_para("⚖   LEGAL NOTICE   ⚖", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, color=(255,255,255), font_name='Arial')
        shade_para(hdr, '1a1a2e')
        sub = add_para("Under the Laws of the Islamic Republic of Pakistan", italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, color=(100,100,100), font_name='Arial')
        sep1 = add_para('', space_after=6); add_border(sep1, 'bottom', sz=18, color='C9A84C')
        t_rd = doc.add_table(rows=1, cols=2); t_rd.style = 'Table Grid'
        c0, c1 = t_rd.rows[0].cells
        c0.text = f"Ref. No: {ref_num}"; c1.text = f"Date: {date_str}"
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for c in [c0, c1]:
            for r in c.paragraphs:
                for run in r.runs: run.font.size = Pt(9)
            no_border(c)
        doc.add_paragraph('')
        t_p = doc.add_table(rows=1, cols=2); t_p.style = 'Table Grid'
        lc, rc = t_p.rows[0].cells
        lp = lc.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = lp.add_run("FROM (Murasil / Noticee):\n"); r1.bold=True; r1.font.size=Pt(9); r1.font.color.rgb=RGBColor(26,26,46)
        r2 = lp.add_run(f"{your_name or 'Applicant'}\n{your_address or 'Pakistan'}"); r2.font.size=Pt(10); r2.bold=True
        rp = rc.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r3 = rp.add_run("TO (Mukhatib / Respondent):\n"); r3.bold=True; r3.font.size=Pt(9); r3.font.color.rgb=RGBColor(139,0,0)
        r4 = rp.add_run(f"{other_name or 'Respondent'}\nAddress as Known to Sender"); r4.font.size=Pt(10); r4.bold=True; r4.font.color.rgb=RGBColor(139,0,0)
        cell_fill(lc, 'eef2ff'); cell_fill(rc, 'fff0f0')
        doc.add_paragraph('')
        mode = add_para("MODE OF SERVICE: REGISTERED POST A.D. / IN PERSON / COURIER", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, color=(80,80,80), font_name='Arial')
        sep2 = add_para('', space_after=8); add_border(sep2, 'bottom', sz=24, color='1a1a2e')
        subj = add_para(f"SUBJECT: {letter_type.upper()}", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8, color=(139,0,0), font_name='Arial')
        shade_para(subj, 'fff5f5')
        sep3 = add_para('', space_before=0, space_after=8); add_border(sep3, 'bottom', sz=4, color='cccccc')
        add_para("Sir / Madam,", bold=True, size=11, space_after=8, font_name='Times New Roman')
        skip = ('sir/', 'madam', 'sincerely', 'yours faithfully', 'yours truly', 'yours sincerely', 'respectfully,', 'regards,')
        for para_text in [p.strip() for p in letter_body.split('\n') if p.strip()]:
            if any(para_text.lower().startswith(x) for x in skip):
                continue
            add_para(para_text, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=9, font_name='Times New Roman')
        sep4 = add_para('', space_after=8); add_border(sep4, 'bottom', sz=4, color='cccccc')
        add_para("Yours faithfully,", size=11, space_after=24, font_name='Times New Roman')
        t_sg = doc.add_table(rows=1, cols=2); t_sg.style = 'Table Grid'
        sc0, sc1 = t_sg.rows[0].cells
        sc0.text = f"________________________\n{your_name or 'Applicant'}\nNoticee / Applicant"
        sc1.text = "________________________\nAdvocate / Legal Counsel\n(If applicable)"
        sc1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        no_border(sc0); no_border(sc1)
        doc.add_paragraph('')
        ack_h = add_para("✦  ACKNOWLEDGEMENT OF RECEIPT  ✦", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, color=(60,60,60), font_name='Arial')
        shade_para(ack_h, 'f3f3f3')
        add_para(
            f"I, {other_name or 'Respondent'}, hereby acknowledge receipt of this Legal Notice "
            f"on _____________ at _____________.\n"
            f"Signature: _________________________     Date: _________________________\n"
            f"Witness: _________________________",
            size=9.5, space_after=4, font_name='Times New Roman', color=(60,60,60))
        doc.add_paragraph('')
        foot = add_para(
            f"Generated by HAQ — Pakistan Legal AI  |  {date_str}  |  Ref: {ref_num}  |  "
            "General information only. Consult a licensed Vakeel for court matters.",
            size=7.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, color=(140,140,140), font_name='Arial')
        add_border(foot, 'top', sz=4, color='cccccc')
        doc.save(path)
        return path
    except Exception as e:
        print(f"DOCX error: {e}")
        return None

def generate_letter(situation, letter_type, your_name, other_name, your_address):
    if not situation or not letter_type:
        return "Please fill in your situation and select letter type.", "", None, None
    prompt = (
        f"Generate a professional Pakistani legal notice body.\n"
        f"Letter Type: {letter_type}\n"
        f"From: {your_name or 'Applicant'}, {your_address or 'Pakistan'}\n"
        f"To: {other_name or 'Respondent'}\n"
        f"Facts: {situation}\n"
        "Write ONLY the 6 body paragraphs. Do NOT include greeting, closing, or signature."
    )
    msgs = [
        {"role": "system", "content": LETTER_SYSTEM_PROMPT},
        {"role": "user",   "content": prompt}
    ]
    body_text, err = call_groq(msgs, max_tokens=1400)
    if not body_text:
        return f"⚠️ {err}", "", None, None
    date_str = datetime.now().strftime("%d %B %Y")
    ref_num = f"LN/{datetime.now().year}/{datetime.now().strftime('%m%d%H%M')}"
    display_text = (
        f"╔══════════════════════════════════════════════════════════╗\n"
        f"║              ⚖  LEGAL NOTICE  ⚖                          ║\n"
        f"║     Under the Laws of Islamic Republic of Pakistan        ║\n"
        f"╚══════════════════════════════════════════════════════════╝\n"
        f"Ref. No: {ref_num}                          Date: {date_str}\n"
        f"FROM:  {your_name or 'Applicant'}\n"
        f"       {your_address or 'Pakistan'}\n"
        f"TO:    {other_name or 'Respondent'}\n"
        f"       [Address as Known to Sender]\n"
        f"MODE:  REGISTERED POST A.D. / IN PERSON\n"
        f"{'─'*60}\n"
        f"SUBJECT: {letter_type.upper()}\n"
        f"{'─'*60}\n"
        f"Sir / Madam,\n"
        f"{body_text}\n"
        f"Yours faithfully,\n"
        f"________________________          ________________________\n"
        f"{your_name or 'Applicant':<30}  Advocate / Legal Counsel\n"
        f"Noticee / Applicant               (If applicable)\n"
        f"{'─'*60}\n"
        f"ACKNOWLEDGEMENT OF RECEIPT\n"
        f"{'─'*60}\n"
        f"I, {other_name or 'Respondent'}, acknowledge receipt on _____________ at _____________.\n"
        f"Signature: _________________________   Date: _________________________\n"
        f"Witness:   _________________________\n"
        f"{'─'*60}\n"
        f"Generated by HAQ — Pakistan Legal AI\n"
        f"Send via Registered Post A.D. Consult a licensed Vakeel for court filings.\n"
        f"{'─'*60}"
    )
    pdf_path = generate_pdf_file(body_text, your_name, other_name, your_address, letter_type)
    docx_path = generate_docx_file(body_text, your_name, other_name, your_address, letter_type)
    verify = build_verify_html(body_text)
    return display_text, verify, pdf_path, docx_path

def save_feedback(history, rating):
    if not history:
        return "No conversation to rate."
    try:
        last_q = history[-1][0] if history else ""
        with open("/tmp/feedback.json", "a") as f:
            f.write(json.dumps({"q": str(last_q)[:100], "r": rating, "t": str(datetime.now())}) + "\n")
    except:
        pass
    return "✓ Thank you!" if rating == "good" else "✓ Noted, we will improve."

# ========== LOCATION-BASED LEGAL HELP DATA ==========
PAKISTAN_CITIES = {
    "Karachi": {
        "province": "Sindh",
        "high_court": "Sindh High Court, Karachi",
        "district_courts": ["City Court Karachi", "Malir District Court", "West District Court", "East District Court", "South District Court"],
        "legal_aid": [
            {"name": "Sindh Legal Empowerment Program (SLEP)", "contact": "021-99204211"},
            {"name": "Legal Aid Society Karachi", "contact": "021-35830041"},
            {"name": "SHRC Legal Aid Cell", "contact": "021-99203351"},
            {"name": "Sindh Bar Council", "contact": "021-99207051"}
        ],
        "police": ["Central Police Office, Karachi", "Sindh Police HQ, Garden"],
        "lawyer_contact": "Sindh Bar Council: 021-99207051 | Karachi Bar: 021-32727232"
    },
    "Hyderabad": {
        "province": "Sindh",
        "high_court": "Sindh High Court Circuit Bench Hyderabad",
        "district_courts": ["District Court Hyderabad", "Sessions Court Hyderabad"],
        "legal_aid": [
            {"name": "Hyderabad District Legal Aid Committee", "contact": "022-9200500"},
            {"name": "Sindh Legal Empowerment Program", "contact": "021-99204211"}
        ],
        "police": ["Hyderabad Police HQ", "Sindh Police Hyderabad Region"],
        "lawyer_contact": "Hyderabad Bar Association: 022-2730271"
    },
    "Sukkur": {
        "province": "Sindh",
        "high_court": "Sindh High Court Circuit Bench Sukkur",
        "district_courts": ["District Court Sukkur", "Sessions Court Sukkur"],
        "legal_aid": [
            {"name": "Sukkur District Legal Aid Committee", "contact": "071-9310363"},
            {"name": "Legal Aid Society Sukkur", "contact": "071-5806444"}
        ],
        "police": ["Sukkur Police HQ", "Sindh Police Sukkur Region"],
        "lawyer_contact": "Sukkur Bar Association: 071-9310363"
    },
    "Larkana": {
        "province": "Sindh",
        "high_court": "Sindh High Court Circuit Bench Larkana",
        "district_courts": ["District Court Larkana"],
        "legal_aid": [
            {"name": "Larkana District Legal Aid Committee", "contact": "074-9410575"}
        ],
        "police": ["Larkana Police HQ"],
        "lawyer_contact": "Larkana Bar Association: 074-9410575"
    },
    "Lahore": {
        "province": "Punjab",
        "high_court": "Lahore High Court, Principal Seat",
        "district_courts": ["District Court Lahore", "Model Town Court", "Cantt Court Lahore", "Sessions Court Lahore"],
        "legal_aid": [
            {"name": "Punjab Legal Empowerment & Justice Dept", "contact": "042-99210012"},
            {"name": "Dastak Legal Aid (Free Helpline)", "contact": "0800-22247"},
            {"name": "AGHS Legal Aid Cell", "contact": "042-35883494"},
            {"name": "Punjab Bar Council", "contact": "042-99210012"}
        ],
        "police": ["Lahore Police HQ, Qila Gujjar Singh", "Punjab Police HQ, Lahore"],
        "lawyer_contact": "Punjab Bar Council: 042-99210012 | Lahore Bar: 042-37350591"
    },
    "Rawalpindi": {
        "province": "Punjab",
        "high_court": "Lahore High Court Rawalpindi Bench",
        "district_courts": ["District Court Rawalpindi", "Sessions Court Rawalpindi"],
        "legal_aid": [
            {"name": "Rawalpindi District Legal Aid Committee", "contact": "051-9290321"},
            {"name": "Dastak Legal Aid", "contact": "0800-22247"}
        ],
        "police": ["Rawalpindi Police HQ", "Punjab Police Rawalpindi Region"],
        "lawyer_contact": "Rawalpindi Bar Association: 051-9290321"
    },
    "Faisalabad": {
        "province": "Punjab",
        "high_court": "Lahore High Court Faisalabad Bench",
        "district_courts": ["District Court Faisalabad"],
        "legal_aid": [
            {"name": "Faisalabad District Legal Aid Committee", "contact": "041-9200300"}
        ],
        "police": ["Faisalabad Police HQ"],
        "lawyer_contact": "Faisalabad Bar Association: 041-9200300"
    },
    "Multan": {
        "province": "Punjab",
        "high_court": "Lahore High Court Multan Bench",
        "district_courts": ["District Court Multan", "Sessions Court Multan"],
        "legal_aid": [
            {"name": "Multan District Legal Aid Committee", "contact": "061-9200300"},
            {"name": "Dastak Legal Aid", "contact": "0800-22247"}
        ],
        "police": ["Multan Police HQ", "Punjab Police Multan Region"],
        "lawyer_contact": "Multan Bar Association: 061-9200300"
    },
    "Gujranwala": {
        "province": "Punjab",
        "high_court": "Lahore High Court Gujranwala Bench",
        "district_courts": ["District Court Gujranwala"],
        "legal_aid": [
            {"name": "Gujranwala District Legal Aid Committee", "contact": "055-9200555"}
        ],
        "police": ["Gujranwala Police HQ"],
        "lawyer_contact": "Gujranwala Bar Association: 055-9200555"
    },
    "Sargodha": {
        "province": "Punjab",
        "high_court": "Lahore High Court Sargodha Bench",
        "district_courts": ["District Court Sargodha"],
        "legal_aid": [
            {"name": "Sargodha District Legal Aid Committee", "contact": "048-9200300"}
        ],
        "police": ["Sargodha Police HQ"],
        "lawyer_contact": "Sargodha Bar Association: 048-9200300"
    },
    "Bahawalpur": {
        "province": "Punjab",
        "high_court": "Lahore High Court Bahawalpur Bench",
        "district_courts": ["District Court Bahawalpur"],
        "legal_aid": [
            {"name": "Bahawalpur District Legal Aid Committee", "contact": "062-9200300"}
        ],
        "police": ["Bahawalpur Police HQ"],
        "lawyer_contact": "Bahawalpur Bar Association: 062-9200300"
    },
    "Peshawar": {
        "province": "Khyber Pakhtunkhwa",
        "high_court": "Peshawar High Court, Principal Seat",
        "district_courts": ["District Court Peshawar", "Sessions Court Peshawar"],
        "legal_aid": [
            {"name": "KPK Legal Aid Services", "contact": "091-9210315"},
            {"name": "Peshawar District Legal Aid Committee", "contact": "091-9210315"},
            {"name": "KPK Bar Council", "contact": "091-9210315"}
        ],
        "police": ["Peshawar Police HQ", "KPK Police HQ"],
        "lawyer_contact": "KPK Bar Council: 091-9210315 | Peshawar Bar: 091-9210315"
    },
    "Mardan": {
        "province": "Khyber Pakhtunkhwa",
        "high_court": "Peshawar High Court Mardan Bench",
        "district_courts": ["District Court Mardan"],
        "legal_aid": [
            {"name": "Mardan District Legal Aid Committee", "contact": "0937-9200300"}
        ],
        "police": ["Mardan Police HQ"],
        "lawyer_contact": "Mardan Bar Association: 0937-9200300"
    },
    "Swat / Mingora": {
        "province": "Khyber Pakhtunkhwa",
        "high_court": "Peshawar High Court Swat Bench",
        "district_courts": ["District Court Swat (Mingora)"],
        "legal_aid": [
            {"name": "Swat District Legal Aid Committee", "contact": "0946-9200300"}
        ],
        "police": ["Swat Police HQ, Mingora"],
        "lawyer_contact": "Swat Bar Association: 0946-9200300"
    },
    "Abbottabad": {
        "province": "Khyber Pakhtunkhwa",
        "high_court": "Peshawar High Court Abbottabad Bench",
        "district_courts": ["District Court Abbottabad"],
        "legal_aid": [
            {"name": "Abbottabad District Legal Aid Committee", "contact": "0992-9200300"}
        ],
        "police": ["Abbottabad Police HQ"],
        "lawyer_contact": "Abbottabad Bar Association: 0992-9200300"
    },
    "Quetta": {
        "province": "Balochistan",
        "high_court": "High Court of Balochistan, Quetta",
        "district_courts": ["District Court Quetta", "Sessions Court Quetta"],
        "legal_aid": [
            {"name": "Balochistan Legal Aid Services", "contact": "081-9200300"},
            {"name": "Balochistan Bar Council", "contact": "081-9200300"}
        ],
        "police": ["Quetta Police HQ", "Balochistan Police HQ"],
        "lawyer_contact": "Balochistan Bar Council: 081-9200300 | Quetta Bar: 081-9200300"
    },
    "Gwadar": {
        "province": "Balochistan",
        "high_court": "High Court of Balochistan Circuit Bench Gwadar",
        "district_courts": ["District Court Gwadar"],
        "legal_aid": [
            {"name": "Gwadar District Legal Aid Committee", "contact": "086-9200300"}
        ],
        "police": ["Gwadar Police HQ"],
        "lawyer_contact": "Gwadar Bar Association: 086-9200300"
    },
    "Islamabad": {
        "province": "Islamabad Capital Territory",
        "high_court": "Islamabad High Court",
        "district_courts": ["District Court Islamabad", "Sessions Court Islamabad", "Family Court Islamabad"],
        "legal_aid": [
            {"name": "Federal Judicial Academy Legal Aid", "contact": "051-9255061"},
            {"name": "Islamabad Bar Council", "contact": "051-2823444"},
            {"name": "AGHS Legal Aid Cell", "contact": "042-35883494"}
        ],
        "police": ["Islamabad Police HQ", "ICT Police"],
        "lawyer_contact": "Islamabad Bar Council: 051-2823444 | District Bar: 051-2823444"
    },
    "Mirpur": {
        "province": "Azad Jammu & Kashmir",
        "high_court": "High Court of Azad Jammu & Kashmir (Mirpur Bench)",
        "district_courts": ["District Court Mirpur"],
        "legal_aid": [
            {"name": "AJK Legal Aid Committee", "contact": "05827-9200300"}
        ],
        "police": ["Mirpur Police HQ"],
        "lawyer_contact": "Mirpur Bar Association: 05827-9200300"
    },
    "Muzaffarabad": {
        "province": "Azad Jammu & Kashmir",
        "high_court": "High Court of Azad Jammu & Kashmir, Muzaffarabad",
        "district_courts": ["District Court Muzaffarabad"],
        "legal_aid": [
            {"name": "AJK Legal Aid Committee", "contact": "05822-9200300"}
        ],
        "police": ["Muzaffarabad Police HQ"],
        "lawyer_contact": "AJK Bar Council: 05822-9200300"
    },
    "Gilgit": {
        "province": "Gilgit-Baltistan",
        "high_court": "Gilgit-Baltistan Chief Court",
        "district_courts": ["District Court Gilgit"],
        "legal_aid": [
            {"name": "Gilgit-Baltistan Legal Aid", "contact": "05811-9200300"}
        ],
        "police": ["Gilgit Police HQ"],
        "lawyer_contact": "Gilgit Bar Association: 05811-9200300"
    }
}

CITY_COORDS = {
    "Karachi": (24.8607, 67.0011),
    "Hyderabad": (25.3960, 68.3578),
    "Sukkur": (27.7052, 68.8574),
    "Larkana": (27.5291, 68.2126),
    "Lahore": (31.5204, 74.3587),
    "Rawalpindi": (33.5651, 73.0169),
    "Faisalabad": (31.4180, 73.0790),
    "Multan": (30.1575, 71.5249),
    "Gujranwala": (32.1877, 74.1945),
    "Sargodha": (32.0836, 72.6711),
    "Bahawalpur": (29.3544, 71.6911),
    "Peshawar": (34.0150, 71.5249),
    "Mardan": (34.1982, 72.0451),
    "Swat / Mingora": (34.7717, 72.3602),
    "Abbottabad": (34.1688, 73.2215),
    "Quetta": (30.1798, 66.9750),
    "Gwadar": (25.2048, 62.3334),
    "Islamabad": (33.6844, 73.0479),
    "Mirpur": (33.1483, 73.7510),
    "Muzaffarabad": (34.3700, 73.4711),
    "Gilgit": (35.8815, 74.4643)
}

PROVINCIAL_LAWS = {
    "Sindh": {
        "tenancy": "Sindh Tenancy Act 1950 — Landlord must give 2 months notice for residential eviction. Eviction only through court. Rent increase max 10% per year.",
        "labour": "Sindh Industrial Relations Act 2013 — Separate labour courts for Sindh. Minimum wage: Rs 37,000/month (2024-25). EOBI & Sindh Social Security applicable.",
        "family": "Sindh Family Courts Act 2013 — Family courts in every district. Khula decree typically in 3-6 months. No husband consent required for khula (MFLO Sec 8).",
        "cyber": "PECA 2016 (Federal) applies. FIA Cyber Crime Wing Karachi handles complaints. Report at: www.ccw.gov.pk or FIA CCW Karachi.",
        "local_gov": "Sindh Local Government Act 2013 — Union Committees handle local disputes, birth/death certificates, municipal issues.",
        "women": "Sindh Domestic Violence (Prevention and Protection) Act 2013 — Protection orders, residence orders, and monetary relief available from court.",
        "info": "Sindh Transparency & Right to Information Act 2016 — File RTI with any Sindh government department. Appeal to Sindh Information Commission."
    },
    "Punjab": {
        "tenancy": "Punjab Rented Premises Act 2009 — Landlord must give 2 months notice. Rent Tribunal for disputes. Rent increase max 10% per year or as per agreement.",
        "labour": "Punjab Industrial Relations Act 2010 — Labour courts in every district. Minimum wage: Rs 37,000/month (2024-25). Punjab Social Security & EOBI.",
        "family": "Punjab Family Courts Act 2014 — Family courts in every district. Khula decree in 3-6 months. Punjab Marriage Registration Rules 2021.",
        "cyber": "PECA 2016 (Federal) applies. FIA Cyber Crime Wing Lahore/Rawalpindi. Punjab Safe Cities Authority for harassment complaints.",
        "local_gov": "Punjab Local Government Act 2022 — Local councils for municipal disputes, property tax, building plans.",
        "women": "Punjab Protection of Women Against Violence Act 2016 — Protection Centres in every district. Toll-free helpline: 1043. Violence Against Women Centres (VAWCs).",
        "info": "Punjab Transparency & Right to Information Act 2013 — File RTI with any Punjab department. Appeal to Punjab Information Commission."
    },
    "Khyber Pakhtunkhwa": {
        "tenancy": "KPK Tenancy Act 1950 — Similar to Sindh. Eviction only through court. 2 months notice required for residential premises.",
        "labour": "KPK Industrial Relations Act 2010 — Separate labour courts. Minimum wage: Rs 37,000/month (2024-25). KPK Employees Social Security Institution.",
        "family": "KPK Family Courts Act 2017 — Family courts in every district. Khula and family disputes. KPK also applies MFLO 1961.",
        "cyber": "PECA 2016 (Federal) applies. FIA Cyber Crime Peshawar. KPK Cyber Crime Unit also active.",
        "local_gov": "KPK Local Government Act 2013 — Village and neighbourhood councils for local disputes and municipal matters.",
        "women": "KPK Domestic Violence Against Women (Prevention and Protection) Act 2021 — Protection orders and shelters.",
        "info": "KPK Right to Information Act 2013 — File RTI with any KPK department. Appeal to KPK Information Commission."
    },
    "Balochistan": {
        "tenancy": "Balochistan Tenancy Act 1948 — Similar provisions. Eviction through court only. 2 months notice standard.",
        "labour": "Balochistan Industrial Relations Act 2010 — Labour courts in Quetta and major cities. Minimum wage: Rs 37,000/month (2024-25).",
        "family": "Balochistan Family Courts Act 2014 — Family courts established in Quetta and major districts.",
        "cyber": "PECA 2016 (Federal) applies. Limited FIA Cyber Crime presence in Quetta. Report via FIA HQ Islamabad.",
        "local_gov": "Balochistan Local Government Act 2010 — Local councils for municipal governance.",
        "women": "Balochistan Domestic Violence (Prevention and Protection) Act 2014 — Protection orders available.",
        "info": "Balochistan Right to Information Act 2016 — File RTI with Balochistan government departments."
    },
    "Islamabad Capital Territory": {
        "tenancy": "Islamabad Rent Restriction Ordinance 2001 — Rent Controller for disputes. 2 months notice for eviction. Rent increase regulated.",
        "labour": "Federal labour laws apply (Industrial Relations Act 2012). Islamabad Labour Court. Minimum wage: Rs 37,000/month (2024-25).",
        "family": "West Pakistan Family Courts Act 1964 — Applies to ICT. Family courts in Islamabad. Also MFLO 1961 applies.",
        "cyber": "PECA 2016. FIA HQ Cyber Crime Wing Islamabad. Direct reporting available at FIA HQ.",
        "local_gov": "ICT Local Government Act 2015 — Metropolitan Corporation Islamabad for local governance.",
        "women": "ICT Domestic Violence (Prevention and Protection) Act 2012 — Protection orders and support centres.",
        "info": "Federal Right of Access to Information Act 2017 — File RTI with federal ministries. Federal Ombudsman for complaints."
    },
    "Azad Jammu & Kashmir": {
        "tenancy": "AJK Tenancy Act applies — Similar to Pakistan tenancy laws. Provincial variations exist.",
        "labour": "AJK Industrial Relations Act — Separate labour courts. Minimum wage aligned with federal rates.",
        "family": "AJK Family Courts Act — Family courts in Muzaffarabad, Mirpur, Rawalakot.",
        "cyber": "PECA 2016 (Federal) applies through AJK Council. Limited local cyber crime infrastructure.",
        "local_gov": "AJK Local Government Act — Local councils for municipal governance.",
        "women": "AJK Domestic Violence Act — Protection mechanisms available through district courts.",
        "info": "AJK Right to Information Act — File RTI with AJK government departments."
    },
    "Gilgit-Baltistan": {
        "tenancy": "Gilgit-Baltistan Tenancy Laws — Local adaptations of tenancy acts. Court-based eviction.",
        "labour": "GB Labour Laws — Apply federal standards with local adaptations.",
        "family": "GB Family Courts — Established in Gilgit and Skardu. MFLO 1961 applies.",
        "cyber": "PECA 2016 (Federal) applies. Limited local FIA presence. Report via Islamabad.",
        "local_gov": "GB Governance Order 2018 — Local governance framework.",
        "women": "GB Women Protection mechanisms through district courts and social welfare.",
        "info": "GB Right to Information framework through local governance structures."
    }
}

def get_nearest_city(lat, lon):
    min_dist = float('inf')
    nearest_city = "Islamabad"
    for city, (c_lat, c_lon) in CITY_COORDS.items():
        dist = ((lat - c_lat)**2 + (lon - c_lon)**2) ** 0.5
        if dist < min_dist:
            min_dist = dist
            nearest_city = city
    is_approx = min_dist > 2.0
    return nearest_city, is_approx

def get_location_help(city_name):
    if not city_name or city_name in ("Auto-detecting...", ""):
        return "Please select your city or click Auto-Detect."
    city = city_name.strip()
    if city not in PAKISTAN_CITIES:
        for c in PAKISTAN_CITIES:
            if city.lower() in c.lower() or c.lower() in city.lower():
                city = c
                break
        else:
            return f"""⚠️ City '{city_name}' not found in database.
Available cities: {', '.join(sorted(PAKISTAN_CITIES.keys()))}
Please select from the dropdown or choose the nearest major city."""
    data = PAKISTAN_CITIES[city]
    prov = data["province"]
    prov_data = PROVINCIAL_LAWS.get(prov, {})
    output = f"""╔══════════════════════════════════════════════════════════════════════╗
║  📍 LOCATION-BASED LEGAL HELP
║  {city.upper()}, {prov.upper()}
╚══════════════════════════════════════════════════════════════════════╝
🏛 COURTS & LEGAL INSTITUTIONS
• High Court: {data['high_court']}
• District Courts: {', '.join(data['district_courts'])}
🚔 POLICE & LAW ENFORCEMENT
• {', '.join(data['police'])}
⚖ FREE LEGAL AID & SUPPORT CENTRES
"""
    for aid in data["legal_aid"]:
        output += f"• {aid['name']}: {aid['contact']}\n"
    output += f"\n📞 LOCAL LAWYER / BAR CONTACTS\n• {data['lawyer_contact']}\n"
    output += f"""
═══════════════════════════════════════════════════════════════════════
📜 PROVINCIAL LAW VARIATIONS — {prov.upper()}
═══════════════════════════════════════════════════════════════════════
🏠 TENANCY / RENT / LANDLORD
{prov_data.get('tenancy', 'N/A')}
💼 LABOUR & EMPLOYMENT
{prov_data.get('labour', 'N/A')}
👪 FAMILY LAW (MARRIAGE, DIVORCE, KHULA)
{prov_data.get('family', 'N/A')}
💻 CYBERCRIME & ONLINE HARASSMENT
{prov_data.get('cyber', 'N/A')}
🏘 LOCAL GOVERNMENT / MUNICIPAL
{prov_data.get('local_gov', 'N/A')}
👩 WOMEN'S PROTECTION & DOMESTIC VIOLENCE
{prov_data.get('women', 'N/A')}
📋 RIGHT TO INFORMATION (RTI)
{prov_data.get('info', 'N/A')}
═══════════════════════════════════════════════════════════════════════
🆘 EMERGENCY & HELPLINE NUMBERS
• Police Emergency: 15
• Rescue / Ambulance: 1122
• Fire Brigade: 16
• Punjab Women Helpline: 1043
• KPK Complaint Cell: 091-9210315
• Sindh Police Helpline: 021-99204211
• FIA Cyber Crime: 1991 (nationwide)
═══════════════════════════════════════════════════════════════════════
💡 KEY TAKEAWAY: Laws vary significantly by province.
   What applies in Punjab may NOT apply in Sindh or KPK.
   Always verify with a local Vakeel (lawyer) before court action.
═══════════════════════════════════════════════════════════════════════"""
    return output

def process_auto_location(loc_str):
    if not loc_str or loc_str == "manual":
        return "Auto-detection failed or denied. Please select your city manually.", ""
    try:
        lat, lon = map(float, loc_str.split(","))
        city, approx = get_nearest_city(lat, lon)
        msg = f"📍 Detected nearest city: {city}"
        if approx:
            msg += " (approximate — verify manually if needed)"
        return msg, get_location_help(city)
    except Exception as e:
        return f"Detection error: {str(e)}. Please select manually.", ""

# ========== CSS ==========
CSS = """
.dark, .dark body, .dark .gradio-container, .dark .wrap, .dark .panel,
.dark .tabitem, .dark .contain, .dark .main, .dark .app {
    background-color: #080A10 !important;
    color: #E8ECF5 !important;
}
.dark *, .dark *::before, .dark *::after {
    color: #E8ECF5 !important;
}
.dark .haq-header, .dark .haq-header * {
    color: inherit !important;
}
.dark .gradio-chatbot .message.user {
    color: #E8ECF5 !important;
}
.dark .gradio-chatbot .message.bot {
    color: #E8ECF5 !important;
}
.dark input, .dark textarea, .dark select,
.dark .gradio-textbox textarea, .dark .gradio-textbox input {
    background-color: #161A24 !important;
    color: #E8ECF5 !important;
    border-color: #252B3A !important;
}
.dark button.primary, .dark .gradio-button.primary {
    background: linear-gradient(135deg, #D4AF37, #B8960C) !important;
    color: #080A10 !important;
}
.dark button.secondary, .dark .gradio-button.secondary {
    background: #161A24 !important;
    color: #8B92A8 !important;
    border-color: #252B3A !important;
}
.dark .gradio-dropdown select, .dark select {
    background-color: #161A24 !important;
    color: #E8ECF5 !important;
    border-color: #252B3A !important;
}
.dark .gradio-radio label {
    background: #161A24 !important;
    border-color: #252B3A !important;
    color: #E8ECF5 !important;
}
.dark .gradio-radio label.selected {
    border-color: #D4AF37 !important;
    background: rgba(212, 175, 55, 0.08) !important;
}
.dark .gr-markdown { color: #8B92A8 !important; }
.dark .gr-markdown h2 { color: #E8ECF5 !important; }
.dark .gr-markdown h3 { color: #D4AF37 !important; }
.dark .gr-markdown strong { color: #e8c97a !important; }
.dark .gr-markdown th { color: #D4AF37 !important; }
.dark .gr-markdown td { color: #8B92A8 !important; }
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&family=Amiri:wght@400;700&display=swap');
:root {
    --gold: #D4AF37;
    --gold-light: #E8C96A;
    --gold-dark: #B8960C;
    --gold-glow: rgba(212, 175, 55, 0.25);
    --dark: #080A10;
    --surface: #0F1218;
    --surface-raised: #161A24;
    --surface-hover: #1E2330;
    --border: #252B3A;
    --border-light: #2E3548;
    --text: #E8ECF5;
    --text-secondary: #8B92A8;
    --text-muted: #5A6278;
    --success: #2ECC71;
    --success-bg: rgba(46, 204, 113, 0.08);
    --warning: #F39C12;
    --error: #E74C3C;
    --radius-sm: 8px;
    --radius-md: 12px;
    --radius-lg: 16px;
    --radius-xl: 24px;
    --shadow-sm: 0 1px 2px rgba(0,0,0,0.3);
    --shadow-md: 0 4px 12px rgba(0,0,0,0.4);
    --shadow-lg: 0 8px 32px rgba(0,0,0,0.5);
    --transition: all 0.25s cubic-bezier(0.4, 0, 0.2, 1);
}
* { box-sizing: border-box; margin: 0; padding: 0; }
html { scroll-behavior: smooth; -webkit-tap-highlight-color: transparent; }
body {
    background: var(--dark) !important;
    color: var(--text) !important;
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif !important;
    min-height: 100vh;
    line-height: 1.6;
    overflow-x: hidden;
}
::-webkit-scrollbar { width: 6px; }
::-webkit-scrollbar-track { background: var(--dark); }
::-webkit-scrollbar-thumb { background: var(--border); border-radius: 3px; }
::-webkit-scrollbar-thumb:hover { background: var(--text-muted); }
.gradio-container {
    max-width: 900px !important;
    margin: 0 auto !important;
    padding: 0 16px !important;
    font-family: 'Inter', sans-serif !important;
    background: var(--dark) !important;
}
.gradio-container .contain {
    padding: 0 !important;
    gap: 0 !important;
}
.gradio-container .wrap {
    background: var(--dark) !important;
    border: none !important;
    padding: 0 !important;
}
.gradio-container .panel {
    background: var(--dark) !important;
    border: none !important;
    padding: 0 !important;
}
.haq-header {
    background: linear-gradient(160deg, var(--surface) 0%, var(--dark) 50%, var(--surface) 100%);
    border-bottom: 2px solid var(--gold);
    padding: 32px 20px 24px;
    text-align: center;
    position: relative;
    overflow: hidden;
    margin: 0 -16px 24px;
}
.haq-header::before {
    content: '';
    position: absolute;
    top: -50%;
    left: 50%;
    transform: translateX(-50%);
    width: 600px;
    height: 600px;
    background: radial-gradient(circle, rgba(212,175,55,0.06) 0%, transparent 70%);
    pointer-events: none;
}
.haq-logo-row {
    display: inline-flex;
    align-items: center;
    gap: 14px;
    margin-bottom: 10px;
    position: relative;
    z-index: 1;
}
.haq-emblem {
    width: 52px;
    height: 52px;
    background: linear-gradient(135deg, var(--gold), var(--gold-dark));
    border-radius: 14px;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 26px;
    box-shadow: 0 6px 24px var(--gold-glow), 0 0 0 1px rgba(212,175,55,0.3);
    flex-shrink: 0;
}
.haq-title {
    font-family: 'Amiri', Georgia, serif;
    font-size: 44px;
    font-weight: 700;
    color: var(--gold);
    letter-spacing: 3px;
    text-shadow: 0 0 30px rgba(212,175,55,0.2);
    line-height: 1;
}
.haq-sub {
    font-size: 11px;
    color: var(--text-muted);
    letter-spacing: 3px;
    text-transform: uppercase;
    margin: 8px 0 18px;
    font-weight: 500;
    position: relative;
    z-index: 1;
}
.haq-stats {
    display: flex;
    justify-content: center;
    max-width: 440px;
    margin: 0 auto 16px;
    background: var(--surface-raised);
    border: 1px solid var(--border);
    border-radius: var(--radius-xl);
    overflow: hidden;
    position: relative;
    z-index: 1;
}
.stat {
    flex: 1;
    padding: 12px 8px;
    text-align: center;
    border-right: 1px solid var(--border);
}
.stat:last-child { border-right: none; }
.stat-num {
    display: block;
    font-size: 18px;
    font-weight: 700;
    color: var(--gold);
    line-height: 1;
}
.stat-lbl {
    display: block;
    font-size: 9px;
    color: var(--text-muted);
    text-transform: uppercase;
    letter-spacing: 0.8px;
    margin-top: 3px;
    font-weight: 500;
}
.haq-pills {
    display: flex;
    flex-wrap: wrap;
    justify-content: center;
    gap: 6px;
    margin-bottom: 14px;
    position: relative;
    z-index: 1;
}
.pill {
    background: rgba(212,175,55,0.08);
    border: 1px solid var(--border);
    color: var(--gold);
    padding: 5px 12px;
    border-radius: 100px;
    font-size: 11px;
    font-weight: 600;
    letter-spacing: 0.3px;
}
.verify-strip {
    background: rgba(46, 204, 113, 0.06);
    border: 1px solid rgba(46, 204, 113, 0.2);
    border-radius: var(--radius-md);
    padding: 8px 16px;
    margin: 0 auto 10px;
    max-width: 520px;
    color: var(--success);
    font-size: 12px;
    font-weight: 600;
    display: flex;
    align-items: center;
    justify-content: center;
    gap: 6px;
    position: relative;
    z-index: 1;
}
.disclaimer {
    background: rgba(243, 156, 18, 0.05);
    border: 1px solid rgba(243, 156, 18, 0.12);
    border-radius: var(--radius-md);
    padding: 10px 16px;
    margin: 0 auto 10px;
    max-width: 520px;
    color: #9a8040;
    font-size: 11px;
    font-weight: 500;
    line-height: 1.5;
    text-align: center;
    position: relative;
    z-index: 1;
}
.dev-line {
    color: var(--text-muted);
    font-size: 11px;
    margin-top: 10px;
    position: relative;
    z-index: 1;
}
.dev-line span {
    color: var(--gold);
    font-weight: 600;
}
.gradio-tabs {
    background: var(--surface) !important;
    border-bottom: 1px solid var(--border) !important;
    border-radius: 0 !important;
    padding: 0 !important;
    margin: 0 -16px 20px !important;
    position: sticky;
    top: 0;
    z-index: 100;
    backdrop-filter: blur(16px);
    -webkit-backdrop-filter: blur(16px);
}
.gradio-tabs .tab-nav {
    display: flex !important;
    gap: 0 !important;
    overflow-x: auto !important;
    scrollbar-width: none !important;
    -ms-overflow-style: none !important;
    padding: 0 16px !important;
    max-width: 900px !important;
    margin: 0 auto !important;
    border: none !important;
    background: transparent !important;
}
.gradio-tabs .tab-nav::-webkit-scrollbar { display: none !important; }
.gradio-tabs .tab-nav button {
    flex: 1 !important;
    min-width: 0 !important;
    padding: 14px 12px !important;
    background: transparent !important;
    border: none !important;
    border-bottom: 2px solid transparent !important;
    border-radius: 0 !important;
    color: var(--text-muted) !important;
    font-size: 13px !important;
    font-weight: 600 !important;
    font-family: 'Inter', sans-serif !important;
    white-space: nowrap !important;
    cursor: pointer !important;
    transition: var(--transition) !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    gap: 8px !important;
}
.gradio-tabs .tab-nav button:hover {
    color: var(--text-secondary) !important;
    background: transparent !important;
}
.gradio-tabs .tab-nav button.selected {
    color: var(--gold) !important;
    border-bottom-color: var(--gold) !important;
    background: transparent !important;
}
.gradio-tabs .tabitem {
    background: var(--dark) !important;
    border: none !important;
    padding: 0 !important;
}
.info-banner {
    background: linear-gradient(135deg, rgba(212,175,55,0.06), rgba(212,175,55,0.02));
    border: 1px solid var(--border);
    border-radius: var(--radius-lg);
    padding: 18px 20px;
    margin-bottom: 20px;
}
.info-banner-title {
    color: var(--gold);
    font-size: 14px;
    font-weight: 700;
    margin-bottom: 6px;
    display: flex;
    align-items: center;
    gap: 8px;
}
.info-banner-text {
    color: var(--text-muted);
    font-size: 12px;
    line-height: 1.7;
}
.gradio-chatbot {
    background: var(--surface) !important;
    border: 1px solid var(--border) !important;
    border-radius: var(--radius-lg) !important;
    overflow: hidden !important;
    margin-bottom: 16px !important;
}
.gradio-chatbot > .bubble-wrap {
    background: var(--surface) !important;
    padding: 16px !important;
}
.gradio-chatbot .message.user {
    background: rgba(212, 175, 55, 0.1) !important;
    border: 1px solid rgba(212, 175, 55, 0.2) !important;
    border-radius: 16px 16px 4px 16px !important;
    color: var(--text) !important;
    padding: 14px 18px !important;
    margin-left: 48px !important;
    margin-bottom: 12px !important;
    line-height: 1.8 !important;
    font-size: 14px !important;
}
.gradio-chatbot .message.bot {
    background: var(--surface-raised) !important;
    border: 1px solid var(--border) !important;
    border-radius: 16px 16px 16px 4px !important;
    color: var(--text) !important;
    padding: 14px 18px !important;
    margin-right: 32px !important;
    margin-bottom: 12px !important;
    line-height: 1.8 !important;
    font-size: 14px !important;
}
.gradio-chatbot .avatar {
    width: 36px !important;
    height: 36px !important;
    border-radius: 50% !important;
}
.gradio-chatbot .avatar.user {
    background: var(--surface-hover) !important;
    border: 1px solid var(--border) !important;
}
.gradio-chatbot .avatar.bot {
    background: linear-gradient(135deg, var(--gold), var(--gold-dark)) !important;
}
.gradio-chatbot .message .time {
    color: var(--text-muted) !important;
    font-size: 11px !important;
    margin-top: 6px !important;
}
input[type="text"], textarea, .gradio-textbox textarea, .gradio-textbox input {
    background: var(--surface-raised) !important;
    border: 1px solid var(--border) !important;
    border-radius: var(--radius-md) !important;
    color: var(--text) !important;
    font-size: 14px !important;
    font-family: 'Inter', sans-serif !important;
    padding: 14px 16px !important;
    line-height: 1.6 !important;
    outline: none !important;
    transition: var(--transition) !important;
}
input[type="text"]:focus, textarea:focus, .gradio-textbox textarea:focus, .gradio-textbox input:focus {
    border-color: var(--gold) !important;
    box-shadow: 0 0 0 3px rgba(212, 175, 55, 0.1) !important;
}
input[type="text"]::placeholder, textarea::placeholder, .gradio-textbox textarea::placeholder, .gradio-textbox input::placeholder {
    color: var(--text-muted) !important;
    font-size: 13px !important;
}
.gradio-textbox label, .gradio-textbox .label, .gradio-dropdown label, .gradio-dropdown .label, .gradio-radio label, .gradio-radio .label {
    color: var(--text-secondary) !important;
    font-size: 11px !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.5px !important;
    margin-bottom: 8px !important;
}
button.primary, .gradio-button.primary, .gradio-button[type="button"].primary {
    background: linear-gradient(135deg, var(--gold), var(--gold-dark)) !important;
    color: var(--dark) !important;
    border: none !important;
    border-radius: var(--radius-md) !important;
    padding: 14px 24px !important;
    font-size: 14px !important;
    font-weight: 700 !important;
    font-family: 'Inter', sans-serif !important;
    cursor: pointer !important;
    transition: var(--transition) !important;
    box-shadow: 0 4px 16px var(--gold-glow) !important;
    width: 100% !important;
}
button.primary:hover, .gradio-button.primary:hover {
    filter: brightness(1.08) !important;
    transform: translateY(-1px) !important;
    box-shadow: 0 6px 24px var(--gold-glow) !important;
}
button.secondary, .gradio-button.secondary {
    background: var(--surface-raised) !important;
    color: var(--text-secondary) !important;
    border: 1px solid var(--border) !important;
    border-radius: var(--radius-md) !important;
    padding: 12px 20px !important;
    font-size: 13px !important;
    font-weight: 600 !important;
    font-family: 'Inter', sans-serif !important;
    cursor: pointer !important;
    transition: var(--transition) !important;
    width: 100% !important;
}
button.secondary:hover, .gradio-button.secondary:hover {
    background: var(--surface-hover) !important;
    border-color: var(--border-light) !important;
    color: var(--text) !important;
}
.gradio-dropdown select, .gradio-dropdown .wrap, select {
    background: var(--surface-raised) !important;
    border: 1px solid var(--border) !important;
    border-radius: var(--radius-md) !important;
    color: var(--text) !important;
    font-size: 14px !important;
    font-family: 'Inter', sans-serif !important;
    padding: 14px 16px !important;
    width: 100% !important;
    outline: none !important;
    transition: var(--transition) !important;
    appearance: none !important;
    -webkit-appearance: none !important;
    background-image: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='12' height='12' viewBox='0 0 12 12'%3E%3Cpath fill='%235A6278' d='M6 8L1 3h10z'/%3E%3C/svg%3E") !important;
    background-repeat: no-repeat !important;
    background-position: right 16px center !important;
    padding-right: 40px !important;
    cursor: pointer !important;
}
.gradio-dropdown select:focus, .gradio-dropdown .wrap:focus, select:focus {
    border-color: var(--gold) !important;
    box-shadow: 0 0 0 3px rgba(212, 175, 55, 0.1) !important;
}
.gradio-radio { display: flex !important; flex-direction: column !important; gap: 8px !important; }
.gradio-radio label {
    background: var(--surface-raised) !important;
    border: 1px solid var(--border) !important;
    border-radius: var(--radius-md) !important;
    padding: 13px 15px !important;
    color: var(--text) !important;
    font-size: 13px !important;
    cursor: pointer !important;
    display: flex !important;
    align-items: center !important;
    gap: 10px !important;
    margin: 0 !important;
    min-height: 50px;
    transition: var(--transition) !important;
}
.gradio-radio label:hover {
    border-color: var(--border-light) !important;
    background: rgba(212, 175, 55, 0.05) !important;
}
.gradio-radio label.selected {
    border-color: var(--gold) !important;
    background: rgba(212, 175, 55, 0.08) !important;
}
.answer-box textarea {
    background: var(--surface) !important;
    border: 1px solid var(--border) !important;
    border-radius: var(--radius-lg) !important;
    color: var(--text) !important;
    font-size: 13px !important;
    line-height: 1.85 !important;
    padding: 20px !important;
    font-family: 'Courier New', monospace !important;
}
.gr-markdown { color: var(--text-secondary) !important; line-height: 1.8 !important; font-size: 14px !important; }
.gr-markdown h2 { color: var(--text) !important; font-size: 22px !important; font-weight: 700 !important; border-bottom: 1px solid var(--border) !important; padding-bottom: 10px !important; margin: 24px 0 14px !important; }
.gr-markdown h3 { color: var(--gold) !important; font-size: 15px !important; font-weight: 600 !important; margin-top: 18px !important; }
.gr-markdown strong { color: #e8c97a !important; }
.gr-markdown table { width: 100% !important; border-collapse: collapse !important; margin: 14px 0 !important; }
.gr-markdown th { background: var(--surface-raised) !important; color: var(--gold) !important; padding: 10px 12px !important; font-size: 11px !important; font-weight: 600 !important; text-transform: uppercase !important; }
.gr-markdown td { padding: 10px 12px !important; border-bottom: 1px solid var(--border) !important; color: var(--text-secondary) !important; font-size: 13px !important; }
.gr-markdown blockquote { border-left: 3px solid var(--gold) !important; padding: 10px 14px !important; margin: 14px 0 !important; background: rgba(212,175,55,0.05) !important; border-radius: 0 var(--radius) var(--radius) 0 !important; color: var(--text-secondary) !important; }
.download-banner {
    background: rgba(212,175,55,0.06);
    border: 1px solid var(--border-light);
    border-radius: var(--radius-lg);
    padding: 14px 18px;
    margin: 16px 0 8px;
}
.download-title {
    color: var(--gold);
    font-size: 13px;
    font-weight: 700;
    margin-bottom: 4px;
}
.gradio-file {
    background: var(--surface-raised) !important;
    border: 1px solid var(--border) !important;
    border-radius: var(--radius-md) !important;
    padding: 16px !important;
}
.gradio-file .file-label {
    color: var(--text) !important;
    font-size: 13px !important;
    font-weight: 600 !important;
}
.gradio-file .file-size {
    color: var(--text-muted) !important;
    font-size: 11px !important;
}
.gradio-row { gap: 12px !important; }
.gradio-column { gap: 12px !important; }
@media (max-width: 768px) {
    .haq-title { font-size: 34px; }
    .haq-stats { max-width: 320px; }
    .gradio-tabs .tab-nav button { padding: 11px 10px !important; font-size: 12px !important; }
    .gradio-chatbot .message.user { margin-left: 20px !important; }
    .gradio-chatbot .message.bot { margin-right: 10px !important; }
    .gradio-row { flex-direction: column !important; }
.voice-status input {
    background: rgba(212, 175, 55, 0.08) !important;
    border: 1px solid rgba(212, 175, 55, 0.3) !important;
    color: #D4AF37 !important;
    text-align: center !important;
    font-size: 12px !important;
    font-weight: 600 !important;
}
.dark .voice-status input {
    background: rgba(212, 175, 55, 0.08) !important;
    border: 1px solid rgba(212, 175, 55, 0.3) !important;
    color: #D4AF37 !important;
}
.tts-audio {
    background: var(--surface-raised) !important;
    border: 1px solid rgba(212, 175, 55, 0.3) !important;
    border-radius: var(--radius-md) !important;
    margin-top: 8px !important;
}
.tts-audio .audio-player {
    background: var(--surface-raised) !important;
}
.gradio-audio {
    background: var(--surface-raised) !important;
    border: 1px solid var(--border) !important;
    border-radius: var(--radius-md) !important;
}
.gradio-audio .record-button {
    background: linear-gradient(135deg, #D4AF37, #B8960C) !important;
    color: #080A10 !important;
    border-radius: 50% !important;
    width: 44px !important;
    height: 44px !important;
}
.gradio-audio .stop-button {
    background: #E74C3C !important;
    color: white !important;
    border-radius: 50% !important;
    width: 44px !important;
    height: 44px !important;
}
}
@media (max-width: 480px) {
    .haq-title { font-size: 28px; }
    .haq-emblem { width: 42px; height: 42px; font-size: 20px; }
    .haq-sub { font-size: 10px; letter-spacing: 2px; }
    .pill { padding: 4px 10px; font-size: 10px; }
    .stat-num { font-size: 16px; }
    .stat-lbl { font-size: 8px; }
}
"""

class DarkTheme(gr.themes.Base):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.name = "haq_dark"
        super().set(
            body_background_fill="#080A10",
            body_background_fill_dark="#080A10",
            body_text_color="#E8ECF5",
            body_text_color_dark="#E8ECF5",
            background_fill_primary="#080A10",
            background_fill_primary_dark="#080A10",
            background_fill_secondary="#0F1218",
            background_fill_secondary_dark="#0F1218",
            block_background_fill="#0F1218",
            block_background_fill_dark="#0F1218",
            block_border_color="#252B3A",
            block_border_color_dark="#252B3A",
            input_background_fill="#161A24",
            input_background_fill_dark="#161A24",
            input_border_color="#252B3A",
            input_border_color_dark="#252B3A",
            button_primary_background_fill="#D4AF37",
            button_primary_background_fill_dark="#D4AF37",
            button_primary_text_color="#080A10",
            button_primary_text_color_dark="#080A10",
            button_secondary_background_fill="#161A24",
            button_secondary_background_fill_dark="#161A24",
            button_secondary_text_color="#8B92A8",
            button_secondary_text_color_dark="#8B92A8",
        )

dark_mode_js = """
function forceDarkMode() {
    document.documentElement.classList.add('dark');
    document.body.classList.add('dark');
    document.querySelector('.gradio-container').classList.add('dark');
    const params = new URLSearchParams(window.location.search);
    if (!params.has('__theme')) {
        params.set('__theme', 'dark');
        window.history.replaceState({}, '', window.location.pathname + '?' + params.toString());
    }
}
forceDarkMode();
"""

# ========== GRADIO UI ==========
with gr.Blocks(title="HAQ - Pakistan Legal AI", theme=DarkTheme(), js=dark_mode_js) as demo:

    gr.HTML("""
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet">
    <div class="haq-header">
        <div class="haq-logo-row"><div class="haq-emblem">&#9878;</div><div class="haq-title">HAQ</div></div>
        <div class="haq-sub">Pakistan Legal AI &nbsp;&middot;&nbsp; Apna Haq Jaano</div>
        <div class="haq-stats">
            <div class="stat"><span class="stat-num">4000+</span><span class="stat-lbl">Laws</span></div>
            <div class="stat"><span class="stat-num">80+</span><span class="stat-lbl">Acts</span></div>
            <div class="stat"><span class="stat-num">2</span><span class="stat-lbl">Languages</span></div>
            <div class="stat"><span class="stat-num">Free</span><span class="stat-lbl">Always</span></div>
        </div>
        <div class="haq-pills">
            <span class="pill">Constitution</span><span class="pill">PPC &amp; CrPC</span>
            <span class="pill">Family Law</span><span class="pill">Labour</span>
            <span class="pill">Property</span><span class="pill">Cybercrime</span>
            <span class="pill">Contract</span><span class="pill">Banking</span>
        </div>
        <div class="verify-strip">&#10003; Every Answer Includes Official Law Verification Links</div>
        <div class="disclaimer">&#9888; General legal information only. For court cases always consult a licensed Vakeel (lawyer).</div>
        <div class="dev-line">Developed by <span>Shahrukh Hussain</span> &mdash; Sukkur, Sindh</div>
    </div>""")

    with gr.Tabs(elem_classes="tab-nav"):

        # TAB 1 — MULTI-TURN CHAT
        with gr.Tab("💬 Ask HAQ"):
            gr.HTML("""
            <div style='height:8px'></div>
            <div class='info-banner'>
                <div class='info-banner-title'>💬 Multi-Turn Legal Conversation</div>
                <div class='info-banner-text'>
                    Ask your question — then ask follow-ups like
                    <b>"What if he doesn't comply?"</b> or <b>"Phir kya hoga?"</b>
                    HAQ remembers the full conversation. Urdu &amp; English both work.
                </div>
            </div>""")
            gr.HTML("""
            <div style="display:inline-flex;align-items:center;gap:6px;background:rgba(16,185,129,0.08);border:1px solid rgba(16,185,129,0.2);border-radius:100px;padding:5px 14px;font-size:11px;font-weight:600;color:#10b981;margin-bottom:12px;">
                <span style="width:8px;height:8px;background:#10b981;border-radius:50%;display:inline-block;"></span>
                HAQ remembers your full conversation — ask follow-ups freely
            </div>""")
            chat_history = gr.State([])
            chatbot = gr.Chatbot(
                value=[],
                label="",
                height=480,
                elem_classes="chatbot-wrap",
                show_label=False,
            )
            with gr.Row():
                chat_input = gr.Textbox(
                    placeholder="Apna legal sawal likhein... e.g. 'Police ne bina warrant arrest kiya' ya 'What is Section 302?'",
                    show_label=False,
                    lines=2,
                    max_lines=5,
                    scale=8,
                )
                send_btn = gr.Button("Send ⚖", variant="primary", scale=1, min_width=100)
            with gr.Row():
                with gr.Column(scale=1):
                    voice_status = gr.Textbox(
                        value="🎙️ Click mic → Speak → Stop → HAQ listens",
                        show_label=False,
                        interactive=False,
                        elem_classes="voice-status"
                    )
                with gr.Column(scale=1):
                    voice_recorder = gr.Audio(
                        sources=["microphone"],
                        type="filepath",
                        label="",
                        show_label=False,
                        editable=False,
                        waveform_options={"waveform_color": "#D4AF37", "waveform_progress_color": "#B8960C"}
                    )
            verify_html = gr.HTML()
            tts_audio = gr.Audio(
                label="🔊 HAQ is speaking...",
                autoplay=True,
                buttons=[],
                interactive=False,
                elem_classes="tts-audio",
                visible=False
            )
            tts_status = gr.Textbox(
                value="",
                show_label=False,
                interactive=False,
                visible=False
            )
            with gr.Row():
                clear_btn = gr.Button("🗑 Clear Chat", variant="secondary")
                tts_btn = gr.Button("🔊 Listen to Answer", variant="secondary")
            gr.HTML("<div style='color:#5A6278;font-size:11px;font-weight:600;text-transform:uppercase;letter-spacing:1px;margin:14px 0 8px;'>Was this helpful?</div>")
            with gr.Row():
                with gr.Column(): good_btn = gr.Button("👍 Yes, helpful!", variant="secondary")
                with gr.Column(): bad_btn = gr.Button("👎 Needs improvement", variant="secondary")
            feedback_out = gr.Textbox(label="", interactive=False, elem_classes="fb-out")
            gr.HTML("<div style='color:#5A6278;font-size:11px;font-weight:600;text-transform:uppercase;letter-spacing:1px;margin:14px 0 8px;'>Try these examples — then ask follow-ups:</div>")
            with gr.Row():
                with gr.Column(): ex1 = gr.Button("🚔 Police arrested me", variant="secondary")
                with gr.Column(): ex2 = gr.Button("⚖ Section 302 kya hai?", variant="secondary")
                with gr.Column(): ex3 = gr.Button("📋 FIR kaise karein?", variant="secondary")
                with gr.Column(): ex4 = gr.Button("💔 Khula ka tarika?", variant="secondary")
            with gr.Row():
                with gr.Column(): ex5 = gr.Button("💼 Salary nahi mili", variant="secondary")
                with gr.Column(): ex6 = gr.Button("🏠 Zameen ka jhagra", variant="secondary")
                with gr.Column(): ex7 = gr.Button("💻 Online harassment?", variant="secondary")
                with gr.Column(): ex8 = gr.Button("📝 Contract broken?", variant="secondary")
            gr.HTML("""
            <div style="margin-top:16px;padding:14px 16px;background:rgba(212,168,76,0.05);
                        border:1px solid rgba(212,168,76,0.15);border-radius:12px;">
                <div style="color:#D4AF37;font-size:12px;font-weight:700;margin-bottom:8px;">
                    💡 Example Multi-Turn Conversation
                </div>
                <div style="color:#5A6278;font-size:12px;line-height:2;">
                    <span style="color:#D4AF37;">You:</span> "Police arrested me without warrant"<br>
                    <span style="color:#10b981;">HAQ:</span> Explains Article 10, CrPC Section 54...<br>
                    <span style="color:#D4AF37;">You:</span> <b>"What if they refuse to release me?"</b><br>
                    <span style="color:#10b981;">HAQ:</span> Explains Habeas Corpus, Article 199... <i>(remembers your case)</i><br>
                    <span style="color:#D4AF37;">You:</span> <b>"Can I sue the police after?"</b><br>
                    <span style="color:#10b981;">HAQ:</span> Explains compensation, which court... <i>(still remembers)</i>
                </div>
            </div>""")

            def tuples_to_messages(history):
                msgs = []
                for user_msg, bot_msg in history:
                    if user_msg:
                        msgs.append({"role": "user", "content": str(user_msg)})
                    if bot_msg:
                        msgs.append({"role": "assistant", "content": str(bot_msg)})
                return msgs

            def respond(message, history):
                history, _ = chat_with_haq(message, history)
                last_answer = history[-1][1] if history else ""
                verify = build_verify_html(last_answer)
                return tuples_to_messages(history), history, "", verify

            def clear_chat():
                return [], [], ""

            send_btn.click(
                respond,
                inputs=[chat_input, chat_history],
                outputs=[chatbot, chat_history, chat_input, verify_html]
            )
            chat_input.submit(
                respond,
                inputs=[chat_input, chat_history],
                outputs=[chatbot, chat_history, chat_input, verify_html]
            )

            def play_tts(history):
                if not history:
                    return None, gr.update(visible=False), "No answer to speak."
                last_answer = history[-1][1] if history[-1][1] else None
                if not last_answer:
                    return None, gr.update(visible=False), "No answer to speak."
                audio_path = text_to_speech(last_answer)
                if audio_path:
                    return audio_path, gr.update(visible=True, value=audio_path), ""
                return None, gr.update(visible=False), "Voice output unavailable. Install gTTS or check text."

            tts_btn.click(
                play_tts,
                inputs=[chat_history],
                outputs=[tts_audio, tts_audio, tts_status]
            )
            clear_btn.click(clear_chat, outputs=[chatbot, chat_history, verify_html])

            def on_voice_recorded(audio_path, history):
                transcribed = transcribe_voice(audio_path)
                if transcribed.startswith("⚠️"):
                    return history, transcribed, transcribed
                history, _ = chat_with_haq(transcribed, history)
                last_answer = history[-1][1] if history else ""
                verify = build_verify_html(last_answer)
                return tuples_to_messages(history), history, "", verify

            voice_recorder.stop_recording(
                on_voice_recorded,
                inputs=[voice_recorder, chat_history],
                outputs=[chatbot, chat_history, chat_input, verify_html]
            )
            voice_recorder.start_recording(
                lambda: "🔴 Recording... Speak your legal question clearly",
                outputs=voice_status
            )
            voice_recorder.stop_recording(
                lambda: "⏳ Transcribing...",
                outputs=voice_status,
                queue=False
            )
            good_btn.click(lambda h: save_feedback(h, "good"), [chat_history], feedback_out)
            bad_btn.click(lambda h: save_feedback(h, "bad"), [chat_history], feedback_out)
            ex1.click(lambda: "Police ne mujhe bina warrant arrest kiya. Mera kya haq hai?", outputs=chat_input)
            ex2.click(lambda: "Section 302 PPC kya hai aur uski saza kya hai?", outputs=chat_input)
            ex3.click(lambda: "FIR kaise darj karein? Police refuse kar rahi hai.", outputs=chat_input)
            ex4.click(lambda: "Main khula lena chahti hun, kya mujhe shohar ki ijazat chahiye?", outputs=chat_input)
            ex5.click(lambda: "Employer ne 3 mahine se salary nahi di. Mera kya haq hai?", outputs=chat_input)
            ex6.click(lambda: "Koi meri zameen pe kabza kar raha hai. Kya karun?", outputs=chat_input)
            ex7.click(lambda: "Koi mujhe online harass kar raha hai. PECA ke tehat kya action le sakta hun?", outputs=chat_input)
            ex8.click(lambda: "Someone broke a contract with me. What can I do legally in Pakistan?", outputs=chat_input)

        # TAB 2 — KNOW YOUR RIGHTS
        with gr.Tab("Know Your Rights"):
            gr.HTML("""
            <div style='height:4px'></div>
            <div class='info-banner'>
                <div class='info-banner-title'>Know Your Legal Rights</div>
                <div class='info-banner-text'>Select your situation — HAQ explains all your rights with exact law citations.</div>
            </div>""")
            situation_r = gr.Radio(choices=[
                "Police arrested me without warrant",
                "Landlord evicted me illegally",
                "Employer did not pay my salary",
                "I am a cybercrime or harassment victim",
                "I want to file for divorce or khula",
                "I have a property dispute",
                "I received a court notice I don't understand",
                "Child custody issue after divorce",
                "Bank is harassing me for loan",
                "I am a victim of workplace harassment"
            ], label="Select Your Situation")
            rights_btn = gr.Button("Tell Me My Legal Rights", variant="primary")
            rights_output = gr.Textbox(lines=16, label="Your Rights and Next Steps", elem_classes="answer-box")
            rights_verify = gr.HTML()
            rights_btn.click(get_rights, inputs=situation_r, outputs=[rights_output, rights_verify])

        # TAB 3 — LEGAL LETTER
        with gr.Tab("Legal Letter"):
            gr.HTML("""
            <div style='height:4px'></div>
            <div class='info-banner'>
                <div class='info-banner-title'>📄 Free Professional Legal Notice Generator</div>
                <div class='info-banner-text'>
                    HAQ generates a <b>professional Pakistani legal notice</b> with exact law citations.
                    Download as <b>PDF</b> or <b>Word (.docx)</b>. Saves Rs 5,000–10,000 in lawyer fees.
                </div>
            </div>""")
            with gr.Row():
                your_name = gr.Textbox(label="Your Full Name", placeholder="Muhammad Ali Khan")
                other_name = gr.Textbox(label="Other Party Name", placeholder="Mr. Ahmed (Landlord)")
            your_address = gr.Textbox(label="Your Address", placeholder="House 123, Street 4, Sukkur, Sindh")
            letter_type = gr.Dropdown(choices=[
                "Legal Notice to Landlord for Illegal Eviction",
                "Legal Notice to Employer for Unpaid Salary",
                "Legal Notice for Breach of Contract",
                "Complaint Letter to Police for FIR Registration",
                "Legal Notice to Online Harasser",
                "Legal Notice for Property Dispute / Land Grabbing",
                "Complaint to Labour Court",
                "Legal Notice for Loan Harassment by Bank",
                "Legal Notice for Recovery of Money",
                "Legal Notice for Defamation",
            ], label="Select Letter Type")
            situation_desc = gr.Textbox(lines=4, label="Describe Your Situation",
                placeholder="What happened? Include dates, amounts, names, and what outcome you want...")
            letter_btn = gr.Button("⚖ Generate Professional Legal Notice", variant="primary")
            letter_output = gr.Textbox(lines=24, label="Your Legal Notice — Professional Format", elem_classes="answer-box")
            letter_verify = gr.HTML()
            gr.HTML("""
            <div class='download-banner'>
                <div class='download-title'>📥 Download Your Legal Notice</div>
            </div>""")
            with gr.Row():
                pdf_download = gr.File(label="📄 Download PDF  (Print-ready)")
                docx_download = gr.File(label="📝 Download Word (.docx)  (Editable)")
            gr.HTML("<div style='color:#5A6278;font-size:11px;text-align:center;margin-top:8px;'>Send via <b>Registered Post A.D.</b> Keep the receipt as legal proof.</div>")
            letter_btn.click(
                generate_letter,
                inputs=[situation_desc, letter_type, your_name, other_name, your_address],
                outputs=[letter_output, letter_verify, pdf_download, docx_download]
            )

        # TAB 4 — DOCUMENT UPLOAD & ANALYSIS
        with gr.Tab("📄 Analyze Document"):
            gr.HTML("""
            <div style='height:4px'></div>
            <div class='info-banner'>
                <div class='info-banner-title'>📄 Document Upload & Legal Analysis</div>
                <div class='info-banner-text'>
                    Upload any legal document — <b>court notices, FIR copies, contracts, agreements, legal notices</b>.
                    HAQ uses <b>OCR (Tesseract)</b> to read the document and <b>AI</b> to explain what it means,
                    highlight critical clauses, and suggest your next steps.
                </div>
            </div>""")
            with gr.Row():
                with gr.Column(scale=1):
                    doc_upload = gr.File(
                        label="Upload Your Document",
                        file_types=[".pdf", ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".txt", ".md"],
                        type="filepath"
                    )
                with gr.Column(scale=1):
                    doc_type = gr.Dropdown(
                        choices=[
                            "Court Notice / Summons",
                            "FIR Copy (First Information Report)",
                            "Contract / Agreement",
                            "Legal Notice Received",
                            "Property Document / Deed",
                            "Employment Letter / Contract",
                            "Rent Agreement / Lease",
                            "Bank / Loan Document",
                            "Divorce / Khula Papers",
                            "Other Legal Document"
                        ],
                        label="What type of document is this?",
                        value="Other Legal Document"
                    )
            analyze_doc_btn = gr.Button("🔍 Analyze My Document", variant="primary")
            doc_status = gr.Textbox(
                label="Processing Status",
                value="Upload a document and select its type, then click Analyze.",
                interactive=False
            )
            doc_analysis_output = gr.Textbox(
                lines=28,
                label="HAQ's Legal Analysis",
                elem_classes="answer-box",
                interactive=False
            )
            doc_extracted_preview = gr.Textbox(
                lines=8,
                label="Extracted Text Preview (OCR)",
                elem_classes="answer-box",
                interactive=False
            )
            doc_verify = gr.HTML()
            gr.HTML("""
            <div style="margin-top:12px;padding:12px 16px;background:rgba(243,156,18,0.05);
                        border:1px solid rgba(243,156,18,0.15);border-radius:12px;">
                <div style="color:#F39C12;font-size:12px;font-weight:700;">
                    ⚠️ Important Notice
                </div>
                <div style="color:#5A6278;font-size:11px;line-height:1.7;margin-top:6px;">
                    • OCR accuracy depends on scan quality. Blurry or handwritten documents may not read well.<br>
                    • For best results, upload clear scanned documents or photos taken in good lighting.<br>
                    • Maximum recommended file size: 10MB. Larger files may timeout.<br>
                    • This analysis is for informational purposes only. Always consult a licensed Vakeel for official legal advice.
                </div>
            </div>""")

            def handle_document_analysis(file_path, doc_type):
                if not file_path:
                    return "Please upload a document first.", "", ""
                status_msg = "⏳ Processing document... extracting text with OCR..."
                analysis, preview = process_uploaded_document(file_path, doc_type)
                verify = build_verify_html(analysis)
                status_msg = "✓ Document analyzed successfully."
                return status_msg, analysis, preview, verify

            analyze_doc_btn.click(
                handle_document_analysis,
                inputs=[doc_upload, doc_type],
                outputs=[doc_status, doc_analysis_output, doc_extracted_preview, doc_verify]
            )

        # TAB 5 — LOCATION-BASED LEGAL HELP
        with gr.Tab("📍 Location Help"):
            gr.HTML("""
            <div style='height:4px'></div>
            <div class='info-banner'>
                <div class='info-banner-title'>📍 Location-Based Legal Help</div>
                <div class='info-banner-text'>
                    Auto-detect your city or select manually. HAQ shows nearest courts, police stations,
                    free legal aid centres, local lawyer contacts, and <b>provincial law variations</b>
                    (Sindh vs Punjab vs KPK vs Balochistan vs ICT).
                </div>
            </div>""")
            with gr.Row():
                with gr.Column(scale=1):
                    auto_detect_btn = gr.Button("📍 Auto-Detect My Location", variant="primary")
                with gr.Column(scale=2):
                    city_dropdown = gr.Dropdown(
                        choices=sorted(list(PAKISTAN_CITIES.keys())),
                        label="Or Select Your City Manually",
                        value="Islamabad"
                    )
            location_hidden = gr.Textbox(visible=False)
            location_status = gr.Textbox(
                label="Detection Status",
                value="Click 'Auto-Detect' or select city manually",
                interactive=False
            )
            location_output = gr.Textbox(
                lines=32,
                label="Your Local Legal Resources & Provincial Laws",
                elem_classes="answer-box",
                interactive=False,
                value=get_location_help("Islamabad")
            )
            gr.HTML("""
            <div style="margin-top:12px;padding:12px 16px;background:rgba(243,156,18,0.05);
                        border:1px solid rgba(243,156,18,0.15);border-radius:12px;">
                <div style="color:#F39C12;font-size:12px;font-weight:700;">
                    ⚠️ Important Notice
                </div>
                <div style="color:#5A6278;font-size:11px;line-height:1.7;margin-top:6px;">
                    • Phone numbers are for government legal aid services and bar councils — not private lawyers.<br>
                    • Provincial laws change frequently. Always verify current rules with a local Vakeel.<br>
                    • For emergencies, call Police <b>15</b>, Rescue <b>1122</b>, or Women Helpline <b>1043</b> (Punjab).
                </div>
            </div>""")
            auto_detect_btn.click(
                None,
                inputs=[],
                outputs=[location_hidden],
                js="""async () => {
                    if (!navigator.geolocation) return 'manual';
                    try {
                        const pos = await new Promise((res, rej) => {
                            navigator.geolocation.getCurrentPosition(res, rej, {timeout: 10000});
                        });
                        return pos.coords.latitude + ',' + pos.coords.longitude;
                    } catch(e) {
                        return 'manual';
                    }
                }"""
            )
            location_hidden.change(
                process_auto_location,
                inputs=[location_hidden],
                outputs=[location_status, location_output]
            )
            city_dropdown.change(
                lambda city: (f"Selected: {city}", get_location_help(city)),
                inputs=[city_dropdown],
                outputs=[location_status, location_output]
            )

        # TAB 6 — ABOUT
        with gr.Tab("About HAQ"):
            gr.HTML("<div style='height:4px'></div>")
            gr.Markdown("""
## What is HAQ?
**HAQ** means "Right" in Urdu. Pakistan's first **RAG-based AI legal assistant** with **multi-turn conversation** — ask follow-up questions and HAQ remembers your full case context.
### Features Overview
| Feature | Description |
|:---|:---|
| **💬 Multi-Turn Chat** | Ask follow-ups, HAQ remembers your full case context |
| **📄 Document Analysis** | Upload court notices, FIRs, contracts — HAQ explains them |
| **⚖ Legal Letter Generator** | Professional legal notices with PDF + DOCX download |
| **📍 Location Help** | Nearest courts, police, legal aid + provincial law variations |
| **🎙️ Voice Input** | Speak your question in Urdu or English |
| **🔊 Voice Output** | Listen to HAQ's answers in your language |
### Multi-Turn Chat — How It Works
| You Say | HAQ Does |
|:---|:---|
| "Police arrested me without warrant" | Explains Article 10, CrPC Section 54 |
| **"What if they refuse to release me?"** | Remembers your case → explains Habeas Corpus |
| **"Can I sue them after?"** | Still remembers → explains compensation, court |
| **"Phir kya hoga?"** | Urdu follow-up → Urdu answer, same context |
HAQ keeps up to **10 turns** of context. Use "Clear Chat" to start a new topic.
### Document Analysis Feature
| Element | How It Works |
|:---|:---|
| **OCR Engine** | Tesseract (free, open-source) — supports English + Urdu |
| **Supported Files** | PDF, PNG, JPG, GIF, BMP, TIFF, TXT, MD |
| **Analysis** | Document overview, critical clauses, deadlines, legal implications, next steps, red flags |
| **Privacy** | Documents are processed in-memory and not stored permanently |
### Legal Letter Feature
| Element | Standard |
|:---|:---|
| Reference Number | LN/YEAR/UNIQUE format |
| Party Structure | FROM (Murasil) / TO (Mukhatib) |
| Legal Citations | Exact law name + section for every claim |
| Demand Deadline | 15-day standard (Pakistani practice) |
| Download | PDF (print-ready) + Word (.docx, editable) |
### Laws Covered
| Category | Coverage |
|:---|:---|
| **Constitutional** | Constitution 1973 + all amendments |
| **Criminal** | PPC 1860, CrPC 1898, Anti-Terrorism Act |
| **Family** | MFLO 1961, DMMA 1939, Family Courts Act |
| **Digital** | PECA 2016, Electronic Transactions Ordinance |
| **Labour** | EOBI, Minimum Wages, Workmen Compensation |
| **Property** | Transfer of Property Act, Registration Act |
| **Commercial** | Contract Act 1872, Banking Laws |
> HAQ provides **general legal information** only.
> For court cases, always consult a licensed Vakeel (lawyer).
**Shahrukh Hussain** — Student, Sukkur, Sindh
*Because 90% of Pakistanis cannot afford a lawyer.*
**"Apna Haq Jaano"** — Know Your Rights
            """)

demo.launch(
    server_name="0.0.0.0",
    server_port=7860,
    show_error=True,
    theme=gr.themes.Base(),
    css=CSS,
)
